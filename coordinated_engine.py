
from __future__ import annotations

import io
import math
import os
import time
import zipfile
from concurrent.futures import ProcessPoolExecutor, as_completed
from itertools import product
from typing import Any, Callable, Dict, Optional, Tuple
from xml.sax.saxutils import escape

import numpy as np
import pandas as pd

from sim_engine import (
    BUS_COUNT,
    ESS_CAP_COL,
    ESS_MAX_COL,
    LOAD_P_COL,
    LOAD_PATTERN_COL,
    LOAD_Q_COL,
    MINUTE_COL,
    PV_P_COL,
    PV_PATTERN_COL,
    TIME_COL,
    WIND_P_COL,
    WIND_PATTERN_COL,
    _get_bus_voltage,
    _normalize_bus_dataframe,
    _safe_runpp,
    create_dynamic_network,
    default_config,
    format_minute,
    prepare_time_profile,
)

SCENARIO_LOAD_INCREASE = "load_increase"
SCENARIO_RENEWABLE_INCREASE = "renewable_increase"
SCENARIO_BOTH_INCREASE = "both_increase"
SCENARIO_RENEWABLE_BY_LOAD_LEVEL = "renewable_by_load_level"
SCENARIO_MODE_HOSTING_CAPACITY = "hosting_capacity"
SCENARIO_MODE_LOAD_PV_MAP = "load_pv_map"
SCENARIO_MODE_ESS_SIZING = "ess_sizing"
LOAD_LEVEL_COL = "부하구간"

STATE_NORMAL = "NORMAL"
STATE_UNDERVOLTAGE = "UNDERVOLTAGE"
STATE_OVERVOLTAGE = "OVERVOLTAGE"
STATE_CONGESTION = "CONGESTION"

CONTROL_CASE_NO_CONTROL = "no_control"
CONTROL_CASE_OLTC_ONLY = "oltc_only"
CONTROL_CASE_OLTC_ESS = "oltc_ess"

SCENARIO_MODE_METADATA: dict[str, dict[str, Any]] = {
    SCENARIO_MODE_HOSTING_CAPACITY: {
        "label": "Hosting Capacity",
        "research_question": "Determine the PV hosting limit under fixed load, ESS, and control assumptions.",
        "fixed_variables": "Load_growth, ESS_size, ESS_location, control_case",
        "varied_variables": "PV_penetration",
    },
    SCENARIO_MODE_LOAD_PV_MAP: {
        "label": "Load-PV Map",
        "research_question": "Build a 2D operating map over joint load and PV stress.",
        "fixed_variables": "ESS_size, ESS_location, control_case",
        "varied_variables": "Load_growth, PV_penetration",
    },
    SCENARIO_MODE_ESS_SIZING: {
        "label": "ESS Sizing",
        "research_question": "Evaluate ESS size and placement under a fixed representative stress case.",
        "fixed_variables": "Load_growth, PV_penetration, control_case",
        "varied_variables": "ESS_size, ESS_location",
    },
}

CONTROL_CASE_LABELS = {
    CONTROL_CASE_NO_CONTROL: "No control",
    CONTROL_CASE_OLTC_ONLY: "OLTC only",
    CONTROL_CASE_OLTC_ESS: "OLTC + ESS",
}


def scenario_label(scenario: str) -> str:
    labels = {
        SCENARIO_LOAD_INCREASE: "부하만 증가",
        SCENARIO_RENEWABLE_INCREASE: "재생에너지 출력 증가",
        SCENARIO_BOTH_INCREASE: "부하와 재생에너지 동시 증가",
        SCENARIO_RENEWABLE_BY_LOAD_LEVEL: "부하구간별 재생에너지 증가",
    }
    return labels.get(scenario, scenario)


def advanced_config(config: Dict[str, float]) -> Dict[str, float]:
    merged = default_config()
    merged.update(config)
    merged.setdefault("voltage_min_limit", 0.94)
    merged.setdefault("voltage_max_limit", 1.06)
    merged.setdefault("voltage_low_off", 0.955)
    merged.setdefault("voltage_high_off", 1.045)
    merged.setdefault("line_limit_mva", 12.0)
    merged.setdefault("line_return_mva", 11.4)
    merged.setdefault("oltc_return_delay_mins", 20)
    merged.setdefault("ess_min_soc", 10.0)
    merged.setdefault("ess_max_soc", 90.0)
    merged.setdefault("ess_p_gain", 16.0)
    merged.setdefault("ess_q_gain", 8.0)
    merged.setdefault("line_relief_gain", 3.0)
    merged.setdefault("ess_ramp_rate_mw_per_min", 0.2)
    merged.setdefault("low_load_upper_pct", 60.0)
    merged.setdefault("mid_load_upper_pct", 85.0)
    merged.setdefault("ess_bus_number", 5)
    merged.setdefault("ess_power_mw", 5.0)
    merged.setdefault("ess_capacity_mwh", 15.0)
    merged.setdefault("control_case", CONTROL_CASE_OLTC_ESS)
    return merged


def prepare_single_ess_bus_df(bus_df: pd.DataFrame, config: Dict[str, float]) -> pd.DataFrame:
    # ESS는 한 대만 사용하므로 입력 데이터와 무관하게 선택한 버스 한 곳에만 용량을 배치한다.
    normalized = _normalize_bus_dataframe(bus_df)
    normalized[ESS_MAX_COL] = 0.0
    normalized[ESS_CAP_COL] = 0.0
    bus_number = int(max(1, min(BUS_COUNT, int(config.get("ess_bus_number", BUS_COUNT)))))
    normalized.at[bus_number - 1, ESS_MAX_COL] = max(0.0, float(config.get("ess_power_mw", 5.0)))
    normalized.at[bus_number - 1, ESS_CAP_COL] = max(0.0, float(config.get("ess_capacity_mwh", 15.0)))
    return normalized


def classify_load_level(load_pct: float, config: Dict[str, float]) -> str:
    if load_pct <= float(config["low_load_upper_pct"]):
        return "경부하"
    if load_pct <= float(config["mid_load_upper_pct"]):
        return "중간부하"
    return "중부하"


def prepare_analysis_profile(
    time_df: pd.DataFrame,
    total_minutes: int,
    time_step_mins: int,
    config: Dict[str, float],
) -> pd.DataFrame:
    prepared = prepare_time_profile(time_df, total_minutes=total_minutes, time_step_mins=time_step_mins).copy()
    prepared[LOAD_LEVEL_COL] = prepared[LOAD_PATTERN_COL].apply(lambda x: classify_load_level(float(x), config))
    return prepared


def evaluate_limits(min_v: float, max_v: float, max_line_mva: float, config: Dict[str, float]) -> Tuple[bool, bool, bool]:
    voltage_ok = min_v >= float(config["voltage_min_limit"]) and max_v <= float(config["voltage_max_limit"])
    line_ok = max_line_mva <= float(config["line_limit_mva"])
    return voltage_ok, line_ok, bool(voltage_ok and line_ok)


def line_metrics(net) -> Tuple[Dict[str, float], float, str, float]:
    if getattr(net, "res_line", pd.DataFrame()).empty:
        return {}, 0.0, "", 0.0
    if "name" not in net.line.columns:
        net.line["name"] = [f"Line {i + 1}" for i in range(len(net.line))]
    else:
        for idx in net.line.index:
            if pd.isna(net.line.at[idx, "name"]) or str(net.line.at[idx, "name"]).strip() == "":
                net.line.at[idx, "name"] = f"Line {int(idx) + 1}"

    line_map: Dict[str, float] = {}
    worst_name = ""
    worst_mva = 0.0
    signed_p = 0.0
    for idx, name in net.line["name"].items():
        p_from = float(net.res_line.at[idx, "p_from_mw"]) if "p_from_mw" in net.res_line.columns else 0.0
        q_from = float(net.res_line.at[idx, "q_from_mvar"]) if "q_from_mvar" in net.res_line.columns else 0.0
        p_to = float(net.res_line.at[idx, "p_to_mw"]) if "p_to_mw" in net.res_line.columns else 0.0
        q_to = float(net.res_line.at[idx, "q_to_mvar"]) if "q_to_mvar" in net.res_line.columns else 0.0
        s_from = math.sqrt(p_from ** 2 + q_from ** 2)
        s_to = math.sqrt(p_to ** 2 + q_to ** 2)
        s_val = max(s_from, s_to)
        line_map[str(name)] = s_val
        if s_val >= worst_mva:
            worst_mva = s_val
            worst_name = str(name)
            signed_p = p_from if abs(p_from) >= abs(p_to) else p_to
    return line_map, worst_mva, worst_name, signed_p


def determine_state(prev_state: str, min_v: float, max_v: float, max_line_mva: float, config: Dict[str, float]) -> str:
    if prev_state == STATE_CONGESTION and max_line_mva > float(config["line_return_mva"]):
        return STATE_CONGESTION
    if prev_state == STATE_UNDERVOLTAGE and min_v < float(config["voltage_low_off"]):
        return STATE_UNDERVOLTAGE
    if prev_state == STATE_OVERVOLTAGE and max_v > float(config["voltage_high_off"]):
        return STATE_OVERVOLTAGE
    if max_line_mva > float(config["line_limit_mva"]):
        return STATE_CONGESTION
    if min_v < float(config["voltage_min_limit"]):
        return STATE_UNDERVOLTAGE
    if max_v > float(config["voltage_max_limit"]):
        return STATE_OVERVOLTAGE
    return STATE_NORMAL


def ramp_to_zero(value: float, ramp_limit: float) -> float:
    if abs(value) <= ramp_limit:
        return 0.0
    return value - math.copysign(ramp_limit, value)


def _distribute_total(total: float, caps: list[float], weights: Optional[list[float]] = None) -> list[float]:
    n = len(caps)
    if n == 0:
        return []
    total = float(total)
    if abs(total) < 1e-9:
        return [0.0] * n
    if weights is None:
        weights = [1.0] * n
    usable = []
    for cap, weight in zip(caps, weights):
        usable.append(max(0.0, float(cap)) * max(0.0, float(weight)))
    denom = sum(usable)
    if denom <= 1e-9:
        return [0.0] * n
    sign = 1.0 if total >= 0.0 else -1.0
    total_abs = abs(total)
    out = []
    for cap, usable_i in zip(caps, usable):
        share = total_abs * usable_i / denom
        out.append(sign * min(float(cap), share))
    return out


def _build_docx_bytes(title: str, paragraphs: list[str]) -> bytes:
    def paragraph_xml(text: str) -> str:
        safe = escape(text)
        return f'<w:p><w:r><w:t xml:space="preserve">{safe}</w:t></w:r></w:p>'

    body = [
        f'<w:p><w:r><w:rPr><w:b/></w:rPr><w:t xml:space="preserve">{escape(title)}</w:t></w:r></w:p>'
    ]
    body.extend(paragraph_xml(p) for p in paragraphs)
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:body>' + ''.join(body) + '<w:sectPr><w:pgSz w:w="12240" w:h="15840"/>'
        '<w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" w:header="720" w:footer="720" w:gutter="0"/>'
        '</w:sectPr></w:body></w:document>'
    )
    content_types = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/></Types>'
    rels = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/></Relationships>'
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/document.xml", document_xml)
    return buffer.getvalue()


def build_analysis_excel_bytes(sim_results: Dict[str, pd.DataFrame]) -> bytes:
    buffer = io.BytesIO()
    last_error = None
    for engine in ["openpyxl", "xlsxwriter", None]:
        try:
            buffer.seek(0)
            buffer.truncate(0)
            writer = pd.ExcelWriter(buffer) if engine is None else pd.ExcelWriter(buffer, engine=engine)
            with writer:
                for sheet_name, key in [
                    ("Voltage_PU", "df_v"),
                    ("OLTC_Tap", "df_tap"),
                    ("ESS_SOC", "df_soc"),
                    ("ESS_P_MW", "df_ess_p"),
                    ("ESS_Q_Mvar", "df_ess_q"),
                    ("Min_Voltage", "df_min_v"),
                    ("Max_Voltage", "df_max_v"),
                    ("State", "df_state"),
                    ("Line_MVA_Max", "df_line_mva_max"),
                    ("Line_MVA", "df_line_mva"),
                    ("Profile", "df_profile"),
                    ("Power_Summary", "df_totals"),
                ]:
                    if key in sim_results and isinstance(sim_results[key], pd.DataFrame):
                        sim_results[key].to_excel(writer, sheet_name=sheet_name)
            return buffer.getvalue()
        except Exception as exc:
            last_error = exc
            continue
    raise RuntimeError("Excel 파일 생성에 실패했습니다.") from last_error



def control_case_label(control_case: str) -> str:
    return CONTROL_CASE_LABELS.get(str(control_case), str(control_case))


def scenario_mode_label(mode: str) -> str:
    meta = SCENARIO_MODE_METADATA.get(str(mode), {})
    return str(meta.get("label", mode))


def scenario_mode_metadata(mode: str) -> Dict[str, Any]:
    mode_key = str(mode)
    if mode_key not in SCENARIO_MODE_METADATA:
        supported = ", ".join(SCENARIO_MODE_METADATA)
        raise ValueError(f"Unsupported scenario mode '{mode_key}'. Supported modes: {supported}")
    meta = dict(SCENARIO_MODE_METADATA[mode_key])
    meta["mode"] = mode_key
    return meta


def supported_scenario_modes() -> list[str]:
    return list(SCENARIO_MODE_METADATA.keys())


def supported_control_cases() -> list[str]:
    return list(CONTROL_CASE_LABELS.keys())


def _stable_float_list(values: list[float]) -> list[float]:
    ordered: list[float] = []
    seen: set[float] = set()
    for value in values:
        key = round(float(value), 10)
        if key not in seen:
            seen.add(key)
            ordered.append(float(value))
    return ordered


def _stable_int_list(values: list[int]) -> list[int]:
    ordered: list[int] = []
    seen: set[int] = set()
    for value in values:
        item = int(value)
        if item not in seen:
            seen.add(item)
            ordered.append(item)
    return ordered


def _stable_text_list(values: list[str]) -> list[str]:
    ordered: list[str] = []
    seen: set[str] = set()
    for value in values:
        text = str(value).strip()
        if text and text not in seen:
            seen.add(text)
            ordered.append(text)
    return ordered


def _expand_scenario_values(spec: Any, default: list[float]) -> list[float]:
    if spec is None:
        return list(default)
    if isinstance(spec, dict):
        if "values" in spec:
            return _stable_float_list([float(x) for x in spec.get("values", [])])
        start = float(spec.get("start", spec.get("min", default[0])))
        stop = float(spec.get("stop", spec.get("max", start)))
        step = float(spec.get("step", 1.0))
        if step <= 0:
            raise ValueError("Scenario range step must be > 0")
        values: list[float] = []
        current = start
        while current <= stop + 1e-9:
            values.append(round(current, 10))
            current += step
        return _stable_float_list(values)
    if isinstance(spec, str):
        text = spec.strip()
        if not text:
            return list(default)
        if ":" in text and "," not in text:
            parts = [part.strip() for part in text.split(":") if part.strip()]
            if len(parts) not in [2, 3]:
                raise ValueError("Range text must be 'start:stop' or 'start:stop:step'")
            start = float(parts[0])
            stop = float(parts[1])
            step = float(parts[2]) if len(parts) == 3 else 1.0
            return _expand_scenario_values({"start": start, "stop": stop, "step": step}, default)
        return _stable_float_list([float(part.strip()) for part in text.split(",") if part.strip()])
    if isinstance(spec, (list, tuple, set, np.ndarray, pd.Series)):
        return _stable_float_list([float(x) for x in spec])
    return [float(spec)]


def _expand_text_values(spec: Any, default: list[str]) -> list[str]:
    if spec is None:
        return list(default)
    if isinstance(spec, str):
        text = spec.strip()
        if not text:
            return list(default)
        return _stable_text_list([part.strip() for part in text.split(",") if part.strip()])
    if isinstance(spec, (list, tuple, set, np.ndarray, pd.Series)):
        return _stable_text_list([str(x).strip() for x in spec])
    return [str(spec).strip()]


def _expand_control_cases(spec: Any, default: list[str]) -> list[str]:
    values = _expand_text_values(spec, default)
    normalized: list[str] = []
    for value in values:
        key = str(value).strip().lower()
        if key not in CONTROL_CASE_LABELS:
            supported = ", ".join(CONTROL_CASE_LABELS)
            raise ValueError(f"Unsupported control_case '{value}'. Supported control cases: {supported}")
        normalized.append(key)
    return _stable_text_list(normalized)


def _expand_location_values(spec: Any, default: list[int]) -> list[int]:
    values = _expand_scenario_values(spec, [float(v) for v in default])
    locations: list[int] = []
    for value in values:
        rounded = int(round(float(value)))
        if abs(float(value) - rounded) > 1e-9:
            raise ValueError("ESS_location must be an integer bus number")
        if rounded < 1 or rounded > BUS_COUNT:
            raise ValueError(f"ESS_location must be between 1 and {BUS_COUNT}")
        locations.append(rounded)
    return _stable_int_list(locations)


def _require_single_value(values: list[Any], field_name: str, mode: str) -> Any:
    if len(values) != 1:
        raise ValueError(f"{scenario_mode_label(mode)} requires a fixed {field_name}, not multiple values")
    return values[0]


def _normalized_execution_key(scenario: Dict[str, Any]) -> tuple[Any, ...]:
    control_case = str(scenario.get("control_case", CONTROL_CASE_OLTC_ESS))
    ess_size = round(float(scenario.get("ESS_size", 0.0)), 10)
    ess_location = int(scenario.get("ESS_location", 0) or 0)
    if control_case != CONTROL_CASE_OLTC_ESS or ess_size <= 0.0:
        ess_size = 0.0
        ess_location = 0
    return (
        str(scenario.get("mode", "")),
        control_case,
        round(float(scenario.get("PV_penetration", 1.0)), 10),
        round(float(scenario.get("Load_growth", 1.0)), 10),
        ess_size,
        ess_location,
        str(scenario.get("base_stress_case", "")),
    )


def _scenario_label_text(mode: str, pv_penetration: float, load_growth: float, ess_size: float, ess_location: int, control_case: str) -> str:
    control_text = control_case_label(control_case)
    if mode == SCENARIO_MODE_HOSTING_CAPACITY:
        return f"HC | PV {pv_penetration:.2f} | Load {load_growth:.2f} | {control_text}"
    if mode == SCENARIO_MODE_LOAD_PV_MAP:
        return f"MAP | Load {load_growth:.2f} | PV {pv_penetration:.2f} | {control_text}"
    return f"ESS | Size {ess_size:.2f} | Bus {ess_location} | PV {pv_penetration:.2f} | Load {load_growth:.2f}"


def _scenario_description_text(mode: str, pv_penetration: float, load_growth: float, ess_size: float, ess_location: int, control_case: str, base_stress_case: str) -> str:
    control_text = control_case_label(control_case)
    if mode == SCENARIO_MODE_HOSTING_CAPACITY:
        return f"Hosting-capacity sweep with PV={pv_penetration:.2f}, load={load_growth:.2f}, control={control_text}."
    if mode == SCENARIO_MODE_LOAD_PV_MAP:
        return f"Operating-map point with load={load_growth:.2f}, PV={pv_penetration:.2f}, control={control_text}."
    return f"ESS sizing case at Bus {ess_location} with size={ess_size:.2f} under {base_stress_case}."


def _make_scenario_dict(
    mode: str,
    pv_penetration: float,
    load_growth: float,
    ess_size: float,
    ess_location: int,
    control_case: str,
    base_stress_case: str = "",
) -> Dict[str, Any]:
    meta = scenario_mode_metadata(mode)
    scenario = {
        "mode": mode,
        "mode_label": meta["label"],
        "research_question": meta["research_question"],
        "PV_penetration": float(pv_penetration),
        "Load_growth": float(load_growth),
        "ESS_size": max(0.0, float(ess_size)),
        "ESS_location": int(ess_location),
        "control_case": str(control_case),
        "control_case_label": control_case_label(str(control_case)),
        "base_stress_case": str(base_stress_case),
    }
    scenario["scenario_label"] = _scenario_label_text(
        mode=mode,
        pv_penetration=float(scenario["PV_penetration"]),
        load_growth=float(scenario["Load_growth"]),
        ess_size=float(scenario["ESS_size"]),
        ess_location=int(scenario["ESS_location"]),
        control_case=str(scenario["control_case"]),
    )
    scenario["scenario_description"] = _scenario_description_text(
        mode=mode,
        pv_penetration=float(scenario["PV_penetration"]),
        load_growth=float(scenario["Load_growth"]),
        ess_size=float(scenario["ESS_size"]),
        ess_location=int(scenario["ESS_location"]),
        control_case=str(scenario["control_case"]),
        base_stress_case=str(scenario["base_stress_case"]),
    )
    return scenario


def _finalize_generated_scenarios(mode: str, scenarios: list[Dict[str, Any]]) -> list[Dict[str, Any]]:
    filtered: list[Dict[str, Any]] = []
    seen: set[tuple[Any, ...]] = set()
    for scenario in scenarios:
        key = _normalized_execution_key(scenario)
        if key in seen:
            continue
        seen.add(key)
        filtered.append(scenario)
    if not filtered:
        raise ValueError(f"{scenario_mode_label(mode)} did not produce any valid scenarios")
    for index, scenario in enumerate(filtered, start=1):
        scenario["scenario_id"] = f"SCN_{index:03d}"
    return filtered


def generate_scenarios(mode: Any, settings: Optional[Dict[str, Any]] = None) -> list[Dict[str, Any]]:
    if settings is None:
        if not isinstance(mode, dict):
            raise ValueError("generate_scenarios requires either (mode, settings) or a settings dict containing 'mode'")
        settings = dict(mode)
        mode = settings.get("mode", SCENARIO_MODE_HOSTING_CAPACITY)
    settings = dict(settings or {})
    mode_key = str(mode or settings.get("mode", SCENARIO_MODE_HOSTING_CAPACITY))
    default_location = int(settings.get("default_ess_location", BUS_COUNT))

    if mode_key == SCENARIO_MODE_HOSTING_CAPACITY:
        pv_values = sorted(_expand_scenario_values(settings.get("pv_penetration", settings.get("PV_penetration")), [1.0]))
        if len(pv_values) < 2:
            raise ValueError("hosting_capacity requires at least two PV penetration values")
        if any(value < 0.0 for value in pv_values):
            raise ValueError("PV_penetration must be >= 0")
        load_growth = float(_require_single_value(_expand_scenario_values(settings.get("load_growth", 1.0), [1.0]), "Load_growth", mode_key))
        ess_size = float(_require_single_value(_expand_scenario_values(settings.get("ess_size", settings.get("ESS_size", 1.0)), [1.0]), "ESS_size", mode_key))
        if load_growth <= 0.0:
            raise ValueError("Load_growth must be > 0")
        if ess_size < 0.0:
            raise ValueError("ESS_size must be >= 0")
        ess_location = int(_require_single_value(_expand_location_values(settings.get("ess_location", default_location), [default_location]), "ESS_location", mode_key))
        control_cases = _expand_control_cases(settings.get("control_case", CONTROL_CASE_OLTC_ESS), [CONTROL_CASE_OLTC_ESS])
        scenarios = []
        for control_case in control_cases:
            effective_ess_size = ess_size if control_case == CONTROL_CASE_OLTC_ESS else 0.0
            for pv_penetration in pv_values:
                scenarios.append(
                    _make_scenario_dict(
                        mode=mode_key,
                        pv_penetration=float(pv_penetration),
                        load_growth=load_growth,
                        ess_size=effective_ess_size,
                        ess_location=ess_location,
                        control_case=control_case,
                    )
                )
        return _finalize_generated_scenarios(mode_key, scenarios)

    if mode_key == SCENARIO_MODE_LOAD_PV_MAP:
        pv_values = sorted(_expand_scenario_values(settings.get("pv_penetration", settings.get("PV_penetration")), [1.0]))
        load_values = sorted(_expand_scenario_values(settings.get("load_growth", settings.get("Load_growth")), [1.0]))
        if len(pv_values) < 2 or len(load_values) < 2:
            raise ValueError("load_pv_map requires at least two PV values and two load values")
        if any(value < 0.0 for value in pv_values):
            raise ValueError("PV_penetration must be >= 0")
        if any(value <= 0.0 for value in load_values):
            raise ValueError("Load_growth must be > 0")
        ess_size = float(_require_single_value(_expand_scenario_values(settings.get("ess_size", settings.get("ESS_size", 1.0)), [1.0]), "ESS_size", mode_key))
        if ess_size < 0.0:
            raise ValueError("ESS_size must be >= 0")
        ess_location = int(_require_single_value(_expand_location_values(settings.get("ess_location", default_location), [default_location]), "ESS_location", mode_key))
        control_case = str(_require_single_value(_expand_control_cases(settings.get("control_case", CONTROL_CASE_OLTC_ESS), [CONTROL_CASE_OLTC_ESS]), "control_case", mode_key))
        effective_ess_size = ess_size if control_case == CONTROL_CASE_OLTC_ESS else 0.0
        scenarios = []
        for load_growth, pv_penetration in product(load_values, pv_values):
            scenarios.append(
                _make_scenario_dict(
                    mode=mode_key,
                    pv_penetration=float(pv_penetration),
                    load_growth=float(load_growth),
                    ess_size=effective_ess_size,
                    ess_location=ess_location,
                    control_case=control_case,
                )
            )
        return _finalize_generated_scenarios(mode_key, scenarios)

    if mode_key == SCENARIO_MODE_ESS_SIZING:
        base_pv_penetration = float(_require_single_value(_expand_scenario_values(settings.get("base_pv_penetration", settings.get("pv_penetration", 1.0)), [1.0]), "PV_penetration", mode_key))
        base_load_growth = float(_require_single_value(_expand_scenario_values(settings.get("base_load_growth", settings.get("load_growth", 1.0)), [1.0]), "Load_growth", mode_key))
        ess_sizes = sorted(_expand_scenario_values(settings.get("ess_size", settings.get("ESS_size", [0.0, 1.0])), [0.0, 1.0]))
        ess_locations = _expand_location_values(settings.get("ess_location", default_location), [default_location])
        if base_pv_penetration < 0.0:
            raise ValueError("PV_penetration must be >= 0")
        if base_load_growth <= 0.0:
            raise ValueError("Load_growth must be > 0")
        if any(value < 0.0 for value in ess_sizes):
            raise ValueError("ESS_size must be >= 0")
        base_ess_power_mw = float(settings.get("base_ess_power_mw", 0.0))
        base_ess_capacity_mwh = float(settings.get("base_ess_capacity_mwh", 0.0))
        if base_ess_power_mw <= 0.0 or base_ess_capacity_mwh <= 0.0:
            raise ValueError("ess_sizing requires a positive base ESS rating in the current configuration")
        control_case = str(_require_single_value(_expand_control_cases(settings.get("control_case", CONTROL_CASE_OLTC_ESS), [CONTROL_CASE_OLTC_ESS]), "control_case", mode_key))
        if control_case != CONTROL_CASE_OLTC_ESS:
            raise ValueError("ess_sizing requires control_case='oltc_ess' because ESS effectiveness is the study target")
        base_stress_case = str(settings.get("base_stress_case") or f"PV {base_pv_penetration:.2f}, Load {base_load_growth:.2f}")
        scenarios = []
        for ess_size, ess_location in product(ess_sizes, ess_locations):
            scenarios.append(
                _make_scenario_dict(
                    mode=mode_key,
                    pv_penetration=base_pv_penetration,
                    load_growth=base_load_growth,
                    ess_size=float(ess_size),
                    ess_location=int(ess_location),
                    control_case=control_case,
                    base_stress_case=base_stress_case,
                )
            )
        finalized = _finalize_generated_scenarios(mode_key, scenarios)
        if len(finalized) < 2:
            raise ValueError("ess_sizing requires at least two distinct ESS size/location scenarios")
        return finalized

    supported = ", ".join(SCENARIO_MODE_METADATA)
    raise ValueError(f"Unsupported scenario mode '{mode_key}'. Supported modes: {supported}")


def build_scenario_preview_df(scenarios: list[Dict[str, Any]]) -> pd.DataFrame:
    if not scenarios:
        return pd.DataFrame()
    rows = []
    for scenario in scenarios:
        rows.append(
            {
                "scenario_id": str(scenario.get("scenario_id", "")),
                "mode": str(scenario.get("mode_label", scenario.get("mode", ""))),
                "control_case": str(scenario.get("control_case_label", scenario.get("control_case", ""))),
                "PV_penetration": float(scenario.get("PV_penetration", np.nan)),
                "Load_growth": float(scenario.get("Load_growth", np.nan)),
                "ESS_size": float(scenario.get("ESS_size", np.nan)),
                "ESS_location": int(scenario.get("ESS_location", 0) or 0),
                "base_stress_case": str(scenario.get("base_stress_case", "")),
                "scenario_label": str(scenario.get("scenario_label", "")),
            }
        )
    preview_df = pd.DataFrame(rows)
    if "base_stress_case" in preview_df.columns and preview_df["base_stress_case"].eq("").all():
        preview_df = preview_df.drop(columns=["base_stress_case"])
    return preview_df



def _scenario_workflow_example_cases(config: Optional[Dict[str, float]] = None) -> Dict[str, list[Dict[str, Any]]]:
    cfg = advanced_config(config or {})
    default_location = int(cfg.get("ess_bus_number", BUS_COUNT))
    example_power = max(float(cfg.get("ess_power_mw", 0.0)), 5.0)
    example_capacity = max(float(cfg.get("ess_capacity_mwh", 0.0)), 15.0)
    shared = {
        "default_ess_location": default_location,
        "base_ess_power_mw": example_power,
        "base_ess_capacity_mwh": example_capacity,
    }
    return {
        SCENARIO_MODE_HOSTING_CAPACITY: generate_scenarios(
            SCENARIO_MODE_HOSTING_CAPACITY,
            {
                **shared,
                "pv_penetration": [0.8, 1.0, 1.2],
                "load_growth": 1.0,
                "ess_size": 1.0,
                "ess_location": default_location,
                "control_case": CONTROL_CASE_OLTC_ESS,
            },
        ),
        SCENARIO_MODE_LOAD_PV_MAP: generate_scenarios(
            SCENARIO_MODE_LOAD_PV_MAP,
            {
                **shared,
                "pv_penetration": [0.8, 1.2],
                "load_growth": [0.9, 1.1],
                "ess_size": 1.0,
                "ess_location": default_location,
                "control_case": CONTROL_CASE_OLTC_ESS,
            },
        ),
        SCENARIO_MODE_ESS_SIZING: generate_scenarios(
            SCENARIO_MODE_ESS_SIZING,
            {
                **shared,
                "base_pv_penetration": 1.6,
                "base_load_growth": 1.0,
                "ess_size": [0.0, 0.5, 1.0],
                "ess_location": default_location,
                "control_case": CONTROL_CASE_OLTC_ESS,
                "base_stress_case": "PV 1.60, Load 1.00",
            },
        ),
    }


def _scenario_flow_example_text(mode: str, scenarios: list[Dict[str, Any]]) -> str:
    if not scenarios:
        return "-"
    if mode == SCENARIO_MODE_HOSTING_CAPACITY:
        return " -> ".join(
            [f"{item['scenario_id']} PV {float(item['PV_penetration']):.2f}" for item in scenarios[:3]]
        )
    if mode == SCENARIO_MODE_LOAD_PV_MAP:
        return ", ".join(
            [
                f"{item['scenario_id']} Load {float(item['Load_growth']):.2f} / PV {float(item['PV_penetration']):.2f}"
                for item in scenarios[:4]
            ]
        )
    return " -> ".join(
        [f"{item['scenario_id']} Size {float(item['ESS_size']):.2f} / Bus {int(item['ESS_location'])}" for item in scenarios[:3]]
    )


def scenario_workflow_lines(config: Optional[Dict[str, float]] = None) -> list[str]:
    examples = _scenario_workflow_example_cases(config)
    return [
        "연구형 배치 시나리오는 무작위 조합이 아니라 연구 질문별 mode로 생성합니다.",
        "공통 흐름: 기준 설정 확정 -> mode별 가변 변수 선택 -> 유효성 검사 및 중복 제거 -> SCN 번호 부여 -> 미리보기 -> 시나리오별 독립 실행 -> 요약 집계",
        "hosting_capacity: 부하, ESS, 제어 조건을 고정한 뒤 PV를 단조 증가시켜 수용 한계와 첫 위반 시점을 찾습니다.",
        f"예시 진행: {_scenario_flow_example_text(SCENARIO_MODE_HOSTING_CAPACITY, examples[SCENARIO_MODE_HOSTING_CAPACITY])}",
        "load_pv_map: ESS와 제어 조건을 고정하고 Load-PV 운전점을 목적성 있게 배치해 2D 운영영역을 만듭니다.",
        f"예시 진행: {_scenario_flow_example_text(SCENARIO_MODE_LOAD_PV_MAP, examples[SCENARIO_MODE_LOAD_PV_MAP])}",
        "ess_sizing: 대표 스트레스 케이스를 고정하고 ESS 크기와 위치를 바꿔 최소 필요 용량과 위치 민감도를 봅니다.",
        f"예시 진행: {_scenario_flow_example_text(SCENARIO_MODE_ESS_SIZING, examples[SCENARIO_MODE_ESS_SIZING])}",
    ]

def _prepare_scenario_inputs(
    config: Dict[str, float],
    bus_df: pd.DataFrame,
    scenario: Optional[Dict[str, Any]],
) -> tuple[Dict[str, float], pd.DataFrame, Dict[str, Any]]:
    scenario_config = advanced_config(config)
    default_location = int(scenario_config.get("ess_bus_number", BUS_COUNT))
    scenario_data = {
        "scenario_id": "SCN_BASE",
        "mode": SCENARIO_MODE_HOSTING_CAPACITY,
        "mode_label": scenario_mode_label(SCENARIO_MODE_HOSTING_CAPACITY),
        "PV_penetration": 1.0,
        "ESS_size": 1.0,
        "Load_growth": 1.0,
        "ESS_location": default_location,
        "control_case": str(scenario_config.get("control_case", CONTROL_CASE_OLTC_ESS)),
        "control_case_label": control_case_label(str(scenario_config.get("control_case", CONTROL_CASE_OLTC_ESS))),
        "base_stress_case": "",
        "scenario_label": "Base case",
        "scenario_description": "Base reference scenario",
    }
    if scenario:
        scenario_data.update(dict(scenario))

    scenario_data["control_case"] = str(scenario_data.get("control_case", CONTROL_CASE_OLTC_ESS)).lower()
    scenario_data["control_case_label"] = control_case_label(str(scenario_data["control_case"]))
    scenario_data["ESS_location"] = int(scenario_data.get("ESS_location", default_location))
    scenario_data["PV_penetration"] = float(scenario_data.get("PV_penetration", 1.0))
    scenario_data["ESS_size"] = max(0.0, float(scenario_data.get("ESS_size", 1.0)))
    scenario_data["Load_growth"] = float(scenario_data.get("Load_growth", 1.0))

    scenario_config["control_case"] = str(scenario_data["control_case"])
    scenario_config["ess_bus_number"] = int(scenario_data["ESS_location"])

    scenario_bus_df = _normalize_bus_dataframe(bus_df).copy()
    scenario_bus_df[PV_P_COL] = pd.to_numeric(scenario_bus_df[PV_P_COL], errors="coerce").fillna(0.0) * float(scenario_data["PV_penetration"])
    scenario_config["ess_power_mw"] = float(scenario_config.get("ess_power_mw", 0.0)) * float(scenario_data["ESS_size"])
    scenario_config["ess_capacity_mwh"] = float(scenario_config.get("ess_capacity_mwh", 0.0)) * float(scenario_data["ESS_size"])
    if str(scenario_config.get("control_case", CONTROL_CASE_OLTC_ESS)) != CONTROL_CASE_OLTC_ESS:
        scenario_config["ess_power_mw"] = 0.0
        scenario_config["ess_capacity_mwh"] = 0.0
    return scenario_config, scenario_bus_df, scenario_data


def _summarize_single_run(
    scenario_data: Dict[str, Any],
    scenario_config: Dict[str, float],
    scenario_bus_df: pd.DataFrame,
    results: Dict[str, pd.DataFrame],
    events: Dict[str, Any],
) -> Dict[str, Any]:
    df_totals = results.get("df_totals", pd.DataFrame())
    voltage_violation_count = int((~df_totals["voltage_ok"]).sum()) if "voltage_ok" in df_totals.columns else int(not bool(events.get("voltage_ok", False)))
    line_violation_count = int((~df_totals["line_ok"]).sum()) if "line_ok" in df_totals.columns else int(not bool(events.get("line_ok", False)))
    overall_violation_count = int((~df_totals["overall_ok"]).sum()) if "overall_ok" in df_totals.columns else int(not bool(events.get("overall_ok", False)))

    return {
        "scenario_id": str(scenario_data.get("scenario_id", "SCN_BASE")),
        "mode": str(scenario_data.get("mode", "")),
        "mode_label": str(scenario_data.get("mode_label", scenario_mode_label(str(scenario_data.get("mode", ""))))),
        "control_case": str(scenario_data.get("control_case", CONTROL_CASE_OLTC_ESS)),
        "control_case_label": str(scenario_data.get("control_case_label", control_case_label(str(scenario_data.get("control_case", CONTROL_CASE_OLTC_ESS))))),
        "PV_penetration": float(scenario_data.get("PV_penetration", 1.0)),
        "ESS_size": float(scenario_data.get("ESS_size", 1.0)),
        "ESS_location": int(scenario_data.get("ESS_location", scenario_config.get("ess_bus_number", BUS_COUNT))),
        "Load_growth": float(scenario_data.get("Load_growth", 1.0)),
        "base_stress_case": str(scenario_data.get("base_stress_case", "")),
        "scenario_label": str(scenario_data.get("scenario_label", "")),
        "scenario_description": str(scenario_data.get("scenario_description", "")),
        "voltage_ok": bool(events.get("voltage_ok", False)),
        "line_ok": bool(events.get("line_ok", False)),
        "overall_ok": bool(events.get("overall_ok", False)),
        "min_voltage": float(events.get("global_min_voltage", np.nan)),
        "max_voltage": float(events.get("global_max_voltage", np.nan)),
        "voltage_violation_count": voltage_violation_count,
        "line_violation_count": line_violation_count,
        "overall_violation_count": overall_violation_count,
        "max_line_loading": float(events.get("global_max_line_mva", np.nan)),
        "OLTC_tap_operation_count": int(events.get("oltc_moves", 0)),
        "ESS_charge_energy": float(events.get("ess_charge_mwh", 0.0)),
        "ESS_discharge_energy": float(events.get("ess_discharge_mwh", 0.0)),
        "convergence_status": str(events.get("convergence_status", "UNKNOWN")),
        "convergence_failures": int(events.get("runpp_failure_count", 0)),
        "violation_flag": bool(not events.get("overall_ok", False)),
        "PV_capacity_total_mw": float(scenario_bus_df[PV_P_COL].sum()),
        "ESS_power_rating_mw": float(scenario_config.get("ess_power_mw", 0.0)),
        "ESS_energy_capacity_mwh": float(scenario_config.get("ess_capacity_mwh", 0.0)),
        "Load_base_total_mw": float(scenario_bus_df[LOAD_P_COL].sum()),
    }

def run_single_simulation(
    config: Dict[str, float],
    bus_df: pd.DataFrame,
    time_df: pd.DataFrame,
    scenario: Optional[Dict[str, Any]] = None,
    ess_efficiency: Optional[float] = None,
    time_step_mins: Optional[int] = None,
    total_minutes: int = 24 * 60,
    include_timeseries: bool = False,
    progress_cb: Optional[Callable[[int, int], None]] = None,
) -> Dict[str, Any]:
    """Run one independent coordinated OLTC-ESS scenario and return summary plus optional detail."""
    scenario_config, scenario_bus_df, scenario_data = _prepare_scenario_inputs(config, bus_df, scenario)
    eff = float(scenario_config.get("ess_efficiency", 0.95) if ess_efficiency is None else ess_efficiency)
    step_mins = int(scenario_config.get("time_step_mins", 10) if time_step_mins is None else time_step_mins)

    results, events = run_coordinated_daily_simulation(
        config=scenario_config,
        bus_df=scenario_bus_df,
        time_df=time_df,
        load_scale=float(scenario_data["Load_growth"]),
        renewable_scale=1.0,
        ess_efficiency=eff,
        total_minutes=total_minutes,
        time_step_mins=step_mins,
        progress_cb=progress_cb,
    )
    summary = _summarize_single_run(scenario_data, scenario_config, scenario_bus_df, results, events)
    output = {
        "scenario": scenario_data,
        "summary": summary,
    }
    if include_timeseries:
        output["results"] = results
        output["events"] = events
    return output


def _batch_worker(payload: Dict[str, Any]) -> Dict[str, Any]:
    scenario = payload.get("scenario", {})
    try:
        return run_single_simulation(
            config=payload["config"],
            bus_df=payload["bus_df"],
            time_df=payload["time_df"],
            scenario=scenario,
            ess_efficiency=payload.get("ess_efficiency"),
            time_step_mins=payload.get("time_step_mins"),
            total_minutes=int(payload.get("total_minutes", 24 * 60)),
            include_timeseries=bool(payload.get("include_timeseries", False)),
            progress_cb=None,
        )
    except Exception as exc:
        return {
            "scenario": scenario,
            "summary": {
                "scenario_id": str(scenario.get("scenario_id", "SCN_ERROR")),
                "mode": str(scenario.get("mode", "")),
                "mode_label": str(scenario.get("mode_label", scenario_mode_label(str(scenario.get("mode", ""))))),
                "control_case": str(scenario.get("control_case", CONTROL_CASE_OLTC_ESS)),
                "control_case_label": str(scenario.get("control_case_label", control_case_label(str(scenario.get("control_case", CONTROL_CASE_OLTC_ESS))))),
                "PV_penetration": float(scenario.get("PV_penetration", 1.0)),
                "ESS_size": float(scenario.get("ESS_size", 1.0)),
                "ESS_location": int(scenario.get("ESS_location", 0) or 0),
                "Load_growth": float(scenario.get("Load_growth", 1.0)),
                "base_stress_case": str(scenario.get("base_stress_case", "")),
                "scenario_label": str(scenario.get("scenario_label", "")),
                "scenario_description": str(scenario.get("scenario_description", "")),
                "min_voltage": np.nan,
                "max_voltage": np.nan,
                "voltage_violation_count": np.nan,
                "line_violation_count": np.nan,
                "overall_violation_count": np.nan,
                "max_line_loading": np.nan,
                "OLTC_tap_operation_count": 0,
                "ESS_charge_energy": 0.0,
                "ESS_discharge_energy": 0.0,
                "convergence_status": "ERROR",
                "convergence_failures": np.nan,
                "violation_flag": True,
                "PV_capacity_total_mw": np.nan,
                "ESS_power_rating_mw": np.nan,
                "ESS_energy_capacity_mwh": np.nan,
                "Load_base_total_mw": np.nan,
                "error_message": str(exc),
            },
        }


def _sort_batch_records(records: list[Dict[str, Any]]) -> list[Dict[str, Any]]:
    return sorted(records, key=lambda item: str(item.get("scenario_id", "")))


def aggregate_batch_results(summary_records: list[Dict[str, Any]]) -> pd.DataFrame:
    """Aggregate high-level batch outcomes for quick comparison and export."""
    summary_df = pd.DataFrame(summary_records)
    if summary_df.empty:
        return pd.DataFrame()
    converged_mask = summary_df["convergence_status"].astype(str).eq("CONVERGED") if "convergence_status" in summary_df.columns else pd.Series(dtype=bool)
    aggregate = {
        "scenario_count": int(len(summary_df)),
        "violation_count": int(summary_df["violation_flag"].fillna(False).sum()) if "violation_flag" in summary_df.columns else 0,
        "converged_count": int(converged_mask.sum()) if not converged_mask.empty else 0,
        "modes": ", ".join(sorted(summary_df["mode_label"].dropna().astype(str).unique())) if "mode_label" in summary_df.columns else "",
        "control_cases": ", ".join(sorted(summary_df["control_case_label"].dropna().astype(str).unique())) if "control_case_label" in summary_df.columns else "",
        "min_of_min_voltage": float(pd.to_numeric(summary_df["min_voltage"], errors="coerce").min()) if "min_voltage" in summary_df.columns else np.nan,
        "max_of_max_voltage": float(pd.to_numeric(summary_df["max_voltage"], errors="coerce").max()) if "max_voltage" in summary_df.columns else np.nan,
        "max_line_loading": float(pd.to_numeric(summary_df["max_line_loading"], errors="coerce").max()) if "max_line_loading" in summary_df.columns else np.nan,
        "max_voltage_violation_count": float(pd.to_numeric(summary_df["voltage_violation_count"], errors="coerce").max()) if "voltage_violation_count" in summary_df.columns else np.nan,
        "max_oltc_operations": float(pd.to_numeric(summary_df["OLTC_tap_operation_count"], errors="coerce").max()) if "OLTC_tap_operation_count" in summary_df.columns else np.nan,
        "total_ess_charge_energy": float(pd.to_numeric(summary_df["ESS_charge_energy"], errors="coerce").sum()) if "ESS_charge_energy" in summary_df.columns else np.nan,
        "total_ess_discharge_energy": float(pd.to_numeric(summary_df["ESS_discharge_energy"], errors="coerce").sum()) if "ESS_discharge_energy" in summary_df.columns else np.nan,
    }
    return pd.DataFrame([aggregate])

def build_batch_summary_csv_bytes(batch_result: Dict[str, Any]) -> bytes:
    """Export batch summary rows to CSV bytes without affecting single-run exports."""
    summary_df = pd.DataFrame(batch_result.get("summary_records", []))
    return summary_df.to_csv(index=False).encode("utf-8-sig")


def build_batch_summary_excel_bytes(batch_result: Dict[str, Any]) -> bytes:
    """Export batch summary and aggregate sheets to Excel bytes."""
    summary_df = pd.DataFrame(batch_result.get("summary_records", []))
    aggregate_df = batch_result.get("aggregate_df", pd.DataFrame())
    buffer = io.BytesIO()
    last_error = None
    for engine in ["openpyxl", "xlsxwriter", None]:
        try:
            buffer.seek(0)
            buffer.truncate(0)
            writer = pd.ExcelWriter(buffer) if engine is None else pd.ExcelWriter(buffer, engine=engine)
            with writer:
                summary_df.to_excel(writer, sheet_name="BatchSummary", index=False)
                if isinstance(aggregate_df, pd.DataFrame) and not aggregate_df.empty:
                    aggregate_df.to_excel(writer, sheet_name="Aggregate", index=False)
            return buffer.getvalue()
        except Exception as exc:
            last_error = exc
            continue
    raise RuntimeError("Batch summary Excel generation failed") from last_error


def run_batch_simulations(
    scenarios: list[Dict[str, Any]],
    config: Dict[str, float],
    bus_df: pd.DataFrame,
    time_df: pd.DataFrame,
    ess_efficiency: Optional[float] = None,
    time_step_mins: Optional[int] = None,
    total_minutes: int = 24 * 60,
    max_workers: Optional[int] = None,
    parallel: bool = True,
    include_timeseries: bool = False,
    progress_cb: Optional[Callable[[int, int, str], None]] = None,
) -> Dict[str, Any]:
    """Execute a scenario batch in serial or with process-level parallelism when safe."""
    scenario_list = list(scenarios)
    parallel_requested = bool(parallel)
    start_time = time.time()
    cpu_count = max(1, int(os.cpu_count() or 1))
    requested_workers = cpu_count if max_workers is None else max(1, int(max_workers))
    worker_count = min(requested_workers, max(1, len(scenario_list)))
    summary_records: list[Dict[str, Any]] = []
    detailed_outputs: list[Dict[str, Any]] = []
    mode = "serial"
    fallback_reason = ""

    if include_timeseries:
        parallel = False
        fallback_reason = "Detailed timeseries collection is forced to serial mode to avoid heavy inter-process transfer."

    def _record_output(result: Dict[str, Any]):
        summary_records.append(result["summary"])
        if include_timeseries:
            detailed_outputs.append(result)

    if parallel and worker_count > 1 and scenario_list:
        payloads = [
            {
                "scenario": scenario,
                "config": config,
                "bus_df": bus_df,
                "time_df": time_df,
                "ess_efficiency": ess_efficiency,
                "time_step_mins": time_step_mins,
                "total_minutes": total_minutes,
                "include_timeseries": False,
            }
            for scenario in scenario_list
        ]
        try:
            with ProcessPoolExecutor(max_workers=worker_count) as executor:
                futures = [executor.submit(_batch_worker, payload) for payload in payloads]
                completed = 0
                for future in as_completed(futures):
                    result = future.result()
                    _record_output(result)
                    completed += 1
                    if progress_cb is not None:
                        progress_cb(completed, len(scenario_list), str(result["summary"].get("scenario_id", completed)))
            mode = "parallel"
        except Exception as exc:
            fallback_reason = str(exc)
            summary_records.clear()
            detailed_outputs.clear()
            parallel = False

    if (not parallel) and scenario_list:
        mode = "serial"
        for idx, scenario in enumerate(scenario_list, start=1):
            result = _batch_worker(
                {
                    "scenario": scenario,
                    "config": config,
                    "bus_df": bus_df,
                    "time_df": time_df,
                    "ess_efficiency": ess_efficiency,
                    "time_step_mins": time_step_mins,
                    "total_minutes": total_minutes,
                    "include_timeseries": include_timeseries,
                }
            )
            _record_output(result)
            if progress_cb is not None:
                progress_cb(idx, len(scenario_list), str(result["summary"].get("scenario_id", idx)))

    summary_records = _sort_batch_records(summary_records)
    aggregate_df = aggregate_batch_results(summary_records)
    elapsed_sec = time.time() - start_time
    result = {
        "summary_records": summary_records,
        "summary_df": pd.DataFrame(summary_records),
        "aggregate_df": aggregate_df,
        "scenarios": scenario_list,
        "execution_mode": mode,
        "parallel_requested": parallel_requested,
        "parallel_used": mode == "parallel",
        "max_workers": worker_count,
        "elapsed_sec": float(elapsed_sec),
        "fallback_reason": fallback_reason,
    }
    if include_timeseries:
        result["detailed_outputs"] = detailed_outputs
    return result

def run_coordinated_daily_simulation(
    config: Dict[str, float],
    bus_df: pd.DataFrame,
    time_df: pd.DataFrame,
    load_scale: float = 1.0,
    renewable_scale: float = 1.0,
    ess_efficiency: float = 0.95,
    total_minutes: int = 24 * 60,
    time_step_mins: int = 10,
    progress_cb: Optional[Callable[[int, int], None]] = None,
) -> Tuple[Dict[str, pd.DataFrame], Dict[str, Any]]:
    # Main coordinated-control loop for one independent time-series run.
    cfg = advanced_config(config)
    control_case = str(cfg.get("control_case", CONTROL_CASE_OLTC_ESS)).lower()
    if control_case not in CONTROL_CASE_LABELS:
        supported = ", ".join(CONTROL_CASE_LABELS)
        raise ValueError(f"Unsupported control_case '{control_case}'. Supported control cases: {supported}")
    cfg["control_case"] = control_case
    oltc_enabled = control_case in [CONTROL_CASE_OLTC_ONLY, CONTROL_CASE_OLTC_ESS]
    ess_enabled = control_case == CONTROL_CASE_OLTC_ESS
    if not ess_enabled:
        cfg["ess_power_mw"] = 0.0
        cfg["ess_capacity_mwh"] = 0.0

    bus_data = prepare_single_ess_bus_df(bus_df, cfg)
    sim_time = prepare_analysis_profile(time_df, total_minutes=total_minutes, time_step_mins=time_step_mins, config=cfg)
    minute_points = sim_time[MINUTE_COL].astype(int).tolist()
    sim_time = sim_time.set_index(MINUTE_COL)
    net = create_dynamic_network(cfg, bus_data)

    ess_efficiency = float(np.clip(ess_efficiency, 1e-6, 1.0))
    load_scale = float(load_scale)
    renewable_scale = float(renewable_scale)
    ramp_rate = float(cfg["ess_ramp_rate_mw_per_min"])

    bus_map = {i + 1: int(net.bus.index[net.bus.name == f"Bus {i + 1}"][0]) for i in range(BUS_COUNT)}
    pv_indices = [int(net.sgen[net.sgen.name == f"PV_{i + 1}"].index[0]) for i in range(BUS_COUNT)]
    wind_indices = [int(net.sgen[net.sgen.name == f"Wind_{i + 1}"].index[0]) for i in range(BUS_COUNT)]
    storage_indices = [int(net.storage[net.storage.name == f"ESS_{i + 1}"].index[0]) for i in range(BUS_COUNT)]
    bus_lookup = {int(idx): str(name) for idx, name in net.bus["name"].items()}

    history_v = {f"Bus {i}": [] for i in range(1, BUS_COUNT + 1)}
    history_soc = {f"Bus {i}": [] for i in range(1, BUS_COUNT + 1)}
    history_ess_p = {f"Bus {i}": [] for i in range(1, BUS_COUNT + 1)}
    history_ess_q = {f"Bus {i}": [] for i in range(1, BUS_COUNT + 1)}
    history_tap = []
    history_min_v = []
    history_max_v = []
    history_state = []
    history_line_max = []
    history_line_map = []
    history_profile = []
    history_totals = []

    current_tap = 0
    current_soc = [float(cfg["ess_init_soc"]) if float(bus_data.at[i, ESS_CAP_COL]) > 0.0 else 0.0 for i in range(BUS_COUNT)]
    current_p_cmd = [0.0] * BUS_COUNT
    current_q_cmd = [0.0] * BUS_COUNT
    current_state = STATE_NORMAL
    oltc_timer_mins = 0.0
    recover_timer_mins = 0.0
    first_violation: Optional[Dict[str, Any]] = None
    global_min_voltage = float("inf")
    global_max_voltage = 0.0
    global_max_line_mva = 0.0
    all_voltage_ok = True
    all_line_ok = True
    ess_charge_mwh = 0.0
    ess_discharge_mwh = 0.0
    runpp_failure_count = 0

    for step_idx, minute in enumerate(minute_points):
        if progress_cb:
            progress_cb(minute, total_minutes)
        delta_mins = float(time_step_mins if step_idx == 0 else max(1, minute - minute_points[step_idx - 1]))
        delta_h = delta_mins / 60.0
        ramp_limit = ramp_rate * delta_mins

        load_pct_raw = float(sim_time.at[minute, LOAD_PATTERN_COL])
        pv_pct_raw = float(sim_time.at[minute, PV_PATTERN_COL])
        wind_pct_raw = float(sim_time.at[minute, WIND_PATTERN_COL])
        load_level = str(sim_time.at[minute, LOAD_LEVEL_COL])
        load_pct = load_pct_raw / 100.0 * load_scale
        pv_pct = pv_pct_raw / 100.0 * renewable_scale
        wind_pct = wind_pct_raw / 100.0 * renewable_scale

        for i in range(BUS_COUNT):
            net.load.at[i, "p_mw"] = float(bus_data.at[i, LOAD_P_COL]) * load_pct
            net.load.at[i, "q_mvar"] = float(bus_data.at[i, LOAD_Q_COL]) * load_pct
            net.sgen.at[pv_indices[i], "p_mw"] = float(bus_data.at[i, PV_P_COL]) * pv_pct
            net.sgen.at[pv_indices[i], "q_mvar"] = 0.0
            net.sgen.at[wind_indices[i], "p_mw"] = float(bus_data.at[i, WIND_P_COL]) * wind_pct
            net.sgen.at[wind_indices[i], "q_mvar"] = 0.0
            net.storage.at[storage_indices[i], "p_mw"] = -current_p_cmd[i]
            net.storage.at[storage_indices[i], "q_mvar"] = -current_q_cmd[i]

        net.trafo.at[0, "tap_pos"] = current_tap
        if not _safe_runpp(net):
            runpp_failure_count += 1

        res_bus = getattr(net, "res_bus", pd.DataFrame())
        if "vm_pu" in res_bus.columns and not res_bus.empty:
            pre_min_v = float(res_bus.vm_pu.min())
            pre_max_v = float(res_bus.vm_pu.max())
            min_v_bus_idx = int(res_bus.vm_pu.idxmin())
            max_v_bus_idx = int(res_bus.vm_pu.idxmax())
        else:
            pre_min_v = 1.0
            pre_max_v = 1.0
            min_v_bus_idx = -1
            max_v_bus_idx = -1
        _, pre_line_mva, _, signed_p = line_metrics(net)
        current_state = determine_state(current_state, pre_min_v, pre_max_v, pre_line_mva, cfg)

        ess_caps = [float(bus_data.at[i, ESS_MAX_COL]) if ess_enabled else 0.0 for i in range(BUS_COUNT)]
        ess_weights = [i + 1 for i in range(BUS_COUNT)]

        if current_state == STATE_CONGESTION:
            if ess_enabled:
                overload = max(0.0, pre_line_mva - float(cfg["line_limit_mva"]))
                p_target = min(sum(ess_caps), overload * float(cfg["line_relief_gain"]))
                if signed_p < 0.0:
                    p_target *= -1.0
                current_p_cmd = _distribute_total(p_target, ess_caps, ess_weights)
                current_q_cmd = [ramp_to_zero(v, ramp_limit) for v in current_q_cmd]
            else:
                current_p_cmd = [0.0] * BUS_COUNT
                current_q_cmd = [0.0] * BUS_COUNT
            if oltc_enabled:
                if pre_max_v > float(cfg["voltage_max_limit"]) and current_tap < 8:
                    oltc_timer_mins += delta_mins
                    if oltc_timer_mins >= float(cfg["oltc_delay_mins"]):
                        current_tap += 1
                        oltc_timer_mins = 0.0
                elif pre_min_v < float(cfg["voltage_min_limit"]) and current_tap > -8:
                    oltc_timer_mins += delta_mins
                    if oltc_timer_mins >= float(cfg["oltc_delay_mins"]):
                        current_tap -= 1
                        oltc_timer_mins = 0.0
                else:
                    oltc_timer_mins = 0.0
            else:
                oltc_timer_mins = 0.0
            recover_timer_mins = 0.0
        elif current_state == STATE_UNDERVOLTAGE:
            if ess_enabled:
                deficit = max(0.0, float(cfg["voltage_min_limit"]) - pre_min_v)
                p_target = min(sum(ess_caps), deficit * float(cfg["ess_p_gain"]))
                q_target = min(sum(ess_caps), deficit * float(cfg["ess_q_gain"]))
                current_p_cmd = _distribute_total(p_target, ess_caps, ess_weights)
                current_q_cmd = _distribute_total(q_target, ess_caps, ess_weights)
            else:
                current_p_cmd = [0.0] * BUS_COUNT
                current_q_cmd = [0.0] * BUS_COUNT
            if oltc_enabled and current_tap > -8:
                oltc_timer_mins += delta_mins
                if oltc_timer_mins >= float(cfg["oltc_delay_mins"]):
                    current_tap -= 1
                    oltc_timer_mins = 0.0
            else:
                oltc_timer_mins = 0.0
            recover_timer_mins = 0.0
        elif current_state == STATE_OVERVOLTAGE:
            if ess_enabled:
                excess = max(0.0, pre_max_v - float(cfg["voltage_max_limit"]))
                p_target = -min(sum(ess_caps), excess * float(cfg["ess_p_gain"]))
                q_target = -min(sum(ess_caps), excess * float(cfg["ess_q_gain"]))
                current_p_cmd = _distribute_total(p_target, ess_caps, ess_weights)
                current_q_cmd = _distribute_total(q_target, ess_caps, ess_weights)
            else:
                current_p_cmd = [0.0] * BUS_COUNT
                current_q_cmd = [0.0] * BUS_COUNT
            if oltc_enabled and current_tap < 8:
                oltc_timer_mins += delta_mins
                if oltc_timer_mins >= float(cfg["oltc_delay_mins"]):
                    current_tap += 1
                    oltc_timer_mins = 0.0
            else:
                oltc_timer_mins = 0.0
            recover_timer_mins = 0.0
        else:
            if ess_enabled:
                current_p_cmd = [ramp_to_zero(v, ramp_limit) for v in current_p_cmd]
                current_q_cmd = [ramp_to_zero(v, ramp_limit) for v in current_q_cmd]
            else:
                current_p_cmd = [0.0] * BUS_COUNT
                current_q_cmd = [0.0] * BUS_COUNT
            if oltc_enabled:
                oltc_timer_mins = 0.0
                recover_timer_mins += delta_mins
                if recover_timer_mins >= float(cfg["oltc_return_delay_mins"]):
                    if current_tap > 0:
                        current_tap -= 1
                        recover_timer_mins = 0.0
                    elif current_tap < 0:
                        current_tap += 1
                        recover_timer_mins = 0.0
            else:
                oltc_timer_mins = 0.0
                recover_timer_mins = 0.0

        for i in range(BUS_COUNT):
            ess_max = float(bus_data.at[i, ESS_MAX_COL])
            ess_cap = float(bus_data.at[i, ESS_CAP_COL])
            soc = current_soc[i]
            if ess_max <= 0.0 or ess_cap <= 0.0:
                current_p_cmd[i] = 0.0
                current_q_cmd[i] = 0.0
                continue
            current_p_cmd[i] = float(np.clip(current_p_cmd[i], -ess_max, ess_max))
            current_q_cmd[i] = float(np.clip(current_q_cmd[i], -ess_max, ess_max))
            if current_p_cmd[i] > 0.0 and soc <= float(cfg["ess_min_soc"]):
                current_p_cmd[i] = 0.0
                current_q_cmd[i] = 0.0
            if current_p_cmd[i] < 0.0 and soc >= float(cfg["ess_max_soc"]):
                current_p_cmd[i] = 0.0
                current_q_cmd[i] = 0.0

        for i in range(BUS_COUNT):
            net.storage.at[storage_indices[i], "p_mw"] = -current_p_cmd[i]
            net.storage.at[storage_indices[i], "q_mvar"] = -current_q_cmd[i]
        net.trafo.at[0, "tap_pos"] = current_tap
        if not _safe_runpp(net):
            runpp_failure_count += 1

        step_charge_mw = 0.0
        step_discharge_mw = 0.0
        for i in range(BUS_COUNT):
            ess_cap = float(bus_data.at[i, ESS_CAP_COL])
            p_cmd = float(current_p_cmd[i])
            if ess_cap <= 0.0 or abs(p_cmd) < 1e-9:
                continue
            if p_cmd > 0.0:
                delta_soc = -(p_cmd * delta_h / ess_cap) * 100.0 / ess_efficiency
                current_soc[i] = float(np.clip(current_soc[i] + delta_soc, 0.0, 100.0))
                step_discharge_mw += p_cmd
            else:
                delta_soc = ((-p_cmd) * delta_h / ess_cap) * 100.0 * ess_efficiency
                current_soc[i] = float(np.clip(current_soc[i] + delta_soc, 0.0, 100.0))
                step_charge_mw += -p_cmd

        ess_charge_mwh += step_charge_mw * delta_h
        ess_discharge_mwh += step_discharge_mw * delta_h

        res_bus = getattr(net, "res_bus", pd.DataFrame())
        if "vm_pu" in res_bus.columns and not res_bus.empty:
            min_v = float(res_bus.vm_pu.min())
            max_v = float(res_bus.vm_pu.max())
            min_v_bus_idx = int(res_bus.vm_pu.idxmin())
            max_v_bus_idx = int(res_bus.vm_pu.idxmax())
        else:
            min_v = 1.0
            max_v = 1.0
            min_v_bus_idx = -1
            max_v_bus_idx = -1
        current_line_map, line_max, line_name, signed_p = line_metrics(net)
        voltage_ok, line_ok, overall_ok = evaluate_limits(min_v, max_v, line_max, cfg)
        all_voltage_ok = all_voltage_ok and voltage_ok
        all_line_ok = all_line_ok and line_ok
        global_min_voltage = min(global_min_voltage, min_v)
        global_max_voltage = max(global_max_voltage, max_v)
        global_max_line_mva = max(global_max_line_mva, line_max)

        total_load = float(net.load["p_mw"].sum())
        total_pv = float(net.sgen.loc[pv_indices, "p_mw"].sum())
        total_wind = float(net.sgen.loc[wind_indices, "p_mw"].sum())
        total_ess = float(sum(-x for x in current_p_cmd))
        net_mw = total_load - total_pv - total_wind - sum(current_p_cmd)

        if first_violation is None and not overall_ok:
            cause = []
            if not voltage_ok:
                cause.append("전압")
            if not line_ok:
                cause.append("선로용량")
            first_violation = {
                "minute": minute,
                "time": format_minute(minute),
                "min_voltage": min_v,
                "max_voltage": max_v,
                "max_line_mva": line_max,
                "line_name": line_name,
                "line_flow_sign_mw": signed_p,
                "min_bus_index": min_v_bus_idx,
                "min_bus_name": bus_lookup.get(min_v_bus_idx, f"BusIndex {min_v_bus_idx}"),
                "max_bus_index": max_v_bus_idx,
                "max_bus_name": bus_lookup.get(max_v_bus_idx, f"BusIndex {max_v_bus_idx}"),
                "cause": "+".join(cause),
                "totals": {
                    "load_mw": total_load,
                    "pv_mw": total_pv,
                    "wind_mw": total_wind,
                    "ess_mw": total_ess,
                    "net_mw": net_mw,
                },
            }

        for i in range(BUS_COUNT):
            history_v[f"Bus {i + 1}"].append(_get_bus_voltage(net, bus_map[i + 1]))
            history_soc[f"Bus {i + 1}"].append(current_soc[i] if float(bus_data.at[i, ESS_CAP_COL]) > 0.0 else np.nan)
            history_ess_p[f"Bus {i + 1}"].append(current_p_cmd[i])
            history_ess_q[f"Bus {i + 1}"].append(current_q_cmd[i])
        history_tap.append(current_tap)
        history_min_v.append(min_v)
        history_max_v.append(max_v)
        history_state.append(current_state)
        history_line_max.append(line_max)
        history_line_map.append(current_line_map)
        history_profile.append(
            {
                TIME_COL: minute / 60.0,
                LOAD_PATTERN_COL: load_pct_raw,
                PV_PATTERN_COL: pv_pct_raw,
                WIND_PATTERN_COL: wind_pct_raw,
                "load_scale": load_scale,
                "renewable_scale": renewable_scale,
                "effective_load_pct": load_pct * 100.0,
                "effective_pv_pct": pv_pct * 100.0,
                "effective_wind_pct": wind_pct * 100.0,
                LOAD_LEVEL_COL: load_level,
            }
        )
        history_totals.append(
            {
                "load_mw": total_load,
                "pv_mw": total_pv,
                "wind_mw": total_wind,
                "ess_mw": total_ess,
                "net_mw": net_mw,
                "voltage_ok": voltage_ok,
                "line_ok": line_ok,
                "overall_ok": overall_ok,
            }
        )

    time_index = [format_minute(m) for m in minute_points[: len(history_tap)]]
    df_v = pd.DataFrame(history_v, index=time_index)
    df_soc = pd.DataFrame(history_soc, index=time_index)
    df_ess_p = pd.DataFrame(history_ess_p, index=time_index)
    df_ess_q = pd.DataFrame(history_ess_q, index=time_index)
    df_tap = pd.DataFrame({"OLTC Tap": history_tap}, index=time_index)
    df_min_v = pd.DataFrame({"Min Voltage (p.u.)": history_min_v}, index=time_index)
    df_max_v = pd.DataFrame({"Max Voltage (p.u.)": history_max_v}, index=time_index)
    df_state = pd.DataFrame({"State": history_state}, index=time_index)
    df_line_mva_max = pd.DataFrame({"Max Line MVA": history_line_max}, index=time_index)
    df_profile = pd.DataFrame(history_profile, index=time_index)
    df_totals = pd.DataFrame(history_totals, index=time_index)
    df_line_mva = pd.DataFrame(history_line_map, index=time_index).fillna(0.0)

    oltc_moves = int(np.count_nonzero(np.diff(np.array(history_tap)))) if len(history_tap) > 1 else 0
    soc_min = float(df_soc.min().min()) if not df_soc.empty else np.nan
    soc_max = float(df_soc.max().max()) if not df_soc.empty else np.nan

    results = {
        "df_v": df_v,
        "df_soc": df_soc,
        "df_ess_p": df_ess_p,
        "df_ess_q": df_ess_q,
        "df_tap": df_tap,
        "df_min_v": df_min_v,
        "df_max_v": df_max_v,
        "df_state": df_state,
        "df_line_mva_max": df_line_mva_max,
        "df_line_mva": df_line_mva,
        "df_profile": df_profile,
        "df_totals": df_totals,
        "time_index": time_index,
    }
    events = {
        "first_violation": first_violation,
        "global_min_voltage": global_min_voltage if np.isfinite(global_min_voltage) else np.nan,
        "global_max_voltage": global_max_voltage if np.isfinite(global_max_voltage) else np.nan,
        "global_max_line_mva": global_max_line_mva if np.isfinite(global_max_line_mva) else np.nan,
        "voltage_ok": bool(all_voltage_ok),
        "line_ok": bool(all_line_ok),
        "overall_ok": bool(all_voltage_ok and all_line_ok),
        "oltc_moves": oltc_moves,
        "final_tap": int(history_tap[-1]) if history_tap else 0,
        "ess_charge_mwh": float(ess_charge_mwh),
        "ess_discharge_mwh": float(ess_discharge_mwh),
        "soc_min": soc_min,
        "soc_max": soc_max,
        "time_step_mins": int(time_step_mins),
        "load_scale": load_scale,
        "renewable_scale": renewable_scale,
        "ess_bus_number": int(cfg["ess_bus_number"]),
        "ess_power_mw": float(cfg["ess_power_mw"]),
        "ess_capacity_mwh": float(cfg["ess_capacity_mwh"]),
        "control_case": control_case,
        "control_case_label": control_case_label(control_case),
        "runpp_failure_count": int(runpp_failure_count),
        "convergence_status": "CONVERGED" if int(runpp_failure_count) == 0 else "FALLBACK_USED",
    }
    return results, events

def _load_level_metrics(results: Dict[str, pd.DataFrame]) -> Dict[str, float]:
    df_profile = results.get("df_profile", pd.DataFrame())
    df_min_v = results.get("df_min_v", pd.DataFrame())
    df_max_v = results.get("df_max_v", pd.DataFrame())
    if df_profile.empty or df_min_v.empty or df_max_v.empty:
        return {}
    joined = pd.concat([df_profile[[LOAD_LEVEL_COL]], df_min_v, df_max_v], axis=1)
    metrics: Dict[str, float] = {}
    for level in ["경부하", "중간부하", "중부하"]:
        part = joined[joined[LOAD_LEVEL_COL] == level]
        metrics[f"{level}_min_voltage"] = float(part["Min Voltage (p.u.)"].min()) if not part.empty else np.nan
        metrics[f"{level}_max_voltage"] = float(part["Max Voltage (p.u.)"].max()) if not part.empty else np.nan
    return metrics


def _run_operating_ranges(results: Dict[str, pd.DataFrame], events: Dict[str, Any]) -> Dict[str, float]:
    # 자동 민감도 분석의 각 실행 단계에서 운전 범위를 요약해 최적화/비교에 활용한다.
    df_totals = results.get("df_totals", pd.DataFrame())
    df_ess_p = results.get("df_ess_p", pd.DataFrame())
    df_tap = results.get("df_tap", pd.DataFrame())
    metrics: Dict[str, float] = {}
    for col, prefix in [("load_mw", "load_total"), ("pv_mw", "pv_total"), ("wind_mw", "wind_total"), ("net_mw", "net_total")]:
        if col in df_totals.columns and not df_totals.empty:
            metrics[f"{prefix}_min_mw"] = float(df_totals[col].min())
            metrics[f"{prefix}_max_mw"] = float(df_totals[col].max())
        else:
            metrics[f"{prefix}_min_mw"] = np.nan
            metrics[f"{prefix}_max_mw"] = np.nan
    if not df_ess_p.empty:
        ess_series = df_ess_p.stack(dropna=True) if hasattr(df_ess_p, "stack") else pd.Series(dtype=float)
        metrics["ess_power_min_mw"] = float(ess_series.min()) if not ess_series.empty else np.nan
        metrics["ess_power_max_mw"] = float(ess_series.max()) if not ess_series.empty else np.nan
    else:
        metrics["ess_power_min_mw"] = np.nan
        metrics["ess_power_max_mw"] = np.nan
    if not df_tap.empty and "OLTC Tap" in df_tap.columns:
        metrics["oltc_tap_min"] = float(df_tap["OLTC Tap"].min())
        metrics["oltc_tap_max"] = float(df_tap["OLTC Tap"].max())
    else:
        metrics["oltc_tap_min"] = np.nan
        metrics["oltc_tap_max"] = np.nan
    metrics["ess_soc_min_pct"] = float(events.get("soc_min", np.nan))
    metrics["ess_soc_max_pct"] = float(events.get("soc_max", np.nan))
    return metrics


def _compact_run_detail(run_info: Dict[str, Any], results: Dict[str, pd.DataFrame], events: Dict[str, Any]) -> Dict[str, Any]:
    # 보고서에서 각 회차별 전압/선로/OLTC/ESS 그래프를 만들 수 있도록 필요한 시계열만 별도로 보관한다.
    compact = dict(run_info)
    compact["results"] = {}
    for key in ["df_min_v", "df_max_v", "df_line_mva_max", "df_tap", "df_ess_p"]:
        df = results.get(key, pd.DataFrame())
        compact["results"][key] = df.copy() if isinstance(df, pd.DataFrame) else pd.DataFrame()
    compact["events"] = {
        "voltage_ok": bool(events.get("voltage_ok", False)),
        "line_ok": bool(events.get("line_ok", False)),
        "overall_ok": bool(events.get("overall_ok", False)),
    }
    return compact


def run_sensitivity_search(
    config: Dict[str, float],
    bus_df: pd.DataFrame,
    time_df: pd.DataFrame,
    scenario: str,
    start_scale: float = 1.0,
    step: float = 0.1,
    max_scale: float = 3.0,
    ess_efficiency: float = 0.95,
    time_step_mins: int = 10,
    total_minutes: int = 24 * 60,
    progress_cb: Optional[Callable[[int, int, float], None]] = None,
) -> Dict[str, Any]:
    cfg = advanced_config(config)
    scale = float(start_scale)
    step = float(step)
    max_scale = float(max_scale)
    runs = []
    run_details = []
    last_results = None
    last_events = None
    first_failure = None
    first_failure_results = None
    first_failure_events = None
    prev_scale = None
    prev_min_voltage = None

    while scale <= max_scale + 1e-9:
        if scenario == SCENARIO_LOAD_INCREASE:
            load_scale = scale
            renewable_scale = 1.0
        elif scenario == SCENARIO_BOTH_INCREASE:
            load_scale = scale
            renewable_scale = scale
        else:
            load_scale = 1.0
            renewable_scale = scale

        def _nested(minute: int, total: int):
            if progress_cb is not None:
                progress_cb(minute, total, scale)

        results, events = run_coordinated_daily_simulation(
            config=cfg,
            bus_df=bus_df,
            time_df=time_df,
            load_scale=load_scale,
            renewable_scale=renewable_scale,
            ess_efficiency=ess_efficiency,
            total_minutes=total_minutes,
            time_step_mins=time_step_mins,
            progress_cb=_nested,
        )
        last_results = results
        last_events = events

        sweep_percent = scale * 100.0
        run_info = {
            "scenario": scenario,
            "sweep_scale": scale,
            "sweep_percent": sweep_percent,
            "load_scale": load_scale,
            "renewable_scale": renewable_scale,
            "min_voltage": float(events.get("global_min_voltage", np.nan)),
            "max_voltage": float(events.get("global_max_voltage", np.nan)),
            "max_line_mva": float(events.get("global_max_line_mva", np.nan)),
            "voltage_ok": bool(events.get("voltage_ok", False)),
            "line_ok": bool(events.get("line_ok", False)),
            "overall_ok": bool(events.get("overall_ok", False)),
            "oltc_moves": int(events.get("oltc_moves", 0)),
            "final_tap": int(events.get("final_tap", 0)),
            "ess_charge_mwh": float(events.get("ess_charge_mwh", 0.0)),
            "ess_discharge_mwh": float(events.get("ess_discharge_mwh", 0.0)),
        }
        run_info.update(_run_operating_ranges(results, events))
        if prev_scale is not None and prev_min_voltage is not None and abs(scale - prev_scale) > 1e-9:
            run_info["min_voltage_sensitivity"] = (run_info["min_voltage"] - prev_min_voltage) / (scale - prev_scale)
        else:
            run_info["min_voltage_sensitivity"] = np.nan
        if scenario == SCENARIO_RENEWABLE_BY_LOAD_LEVEL:
            run_info.update(_load_level_metrics(results))
        if first_failure is None and not run_info["overall_ok"]:
            first_failure = run_info.copy()
            first_failure_results = results
            first_failure_events = events
            if events.get("first_violation"):
                first_failure["first_violation"] = events["first_violation"]
        runs.append(run_info)
        run_details.append(_compact_run_detail(run_info, results, events))
        prev_scale = scale
        prev_min_voltage = run_info["min_voltage"]
        scale = round(scale + step, 10)

    return {
        "scenario": scenario,
        "scenario_label": scenario_label(scenario),
        "start_scale": float(start_scale),
        "step": float(step),
        "max_scale": float(max_scale),
        "time_step_mins": int(time_step_mins),
        "limits": {
            "voltage_min": float(cfg["voltage_min_limit"]),
            "voltage_max": float(cfg["voltage_max_limit"]),
            "line_limit_mva": float(cfg["line_limit_mva"]),
        },
        "runs": runs,
        "run_details": run_details,
        "first_failure": first_failure,
        "first_failure_results": first_failure_results,
        "first_failure_events": first_failure_events,
        "last_results": last_results,
        "last_events": last_events,
        "config": cfg,
        "bus_df": prepare_single_ess_bus_df(bus_df, cfg),
    }


def _report_lines(search_result: Dict[str, Any]) -> list[str]:
    # 보고서 본문은 설정, 판정 결과, 진단 메시지를 한 번에 재구성한다.
    scenario = str(search_result.get("scenario_label", search_result.get("scenario", "")))
    limits = search_result.get("limits", {})
    runs = search_result.get("runs", [])
    first_failure = search_result.get("first_failure")
    last_events = search_result.get("last_events") or {}
    run_details = search_result.get("run_details", [])
    bus_df = search_result.get("bus_df", pd.DataFrame())

    lines = [
        f"시뮬레이션 유형: {scenario}",
        f"증가 시작값: {float(search_result.get('start_scale', 1.0)):.2f}",
        f"증가 간격: {float(search_result.get('step', 0.1)):.2f}",
        f"최대 증가값: {float(search_result.get('max_scale', 3.0)):.2f}",
        f"시뮬레이션 간격: {int(search_result.get('time_step_mins', 10))}분",
        f"전압 허용범위: {float(limits.get('voltage_min', 0.94)):.3f} ~ {float(limits.get('voltage_max', 1.06)):.3f} p.u.",
        f"선로용량 허용치: {float(limits.get('line_limit_mva', 12.0)):.2f} MVA",
        "",
        "연구형 시나리오 생성 및 진행 방식",
    ]
    lines.extend(scenario_workflow_lines(search_result.get("config", {})))
    lines.extend(["", "계통 기본 구성"])
    if isinstance(bus_df, pd.DataFrame) and not bus_df.empty:
        for _, row in bus_df.iterrows():
            lines.append(
                f"{row.iloc[0]}: Load {float(row[LOAD_P_COL]):.2f} MW / PV {float(row[PV_P_COL]):.2f} MW / Wind {float(row[WIND_P_COL]):.2f} MW / ESS {float(row[ESS_MAX_COL]):.2f} MW"
            )
    lines.append("")
    lines.append("현재 운영 알고리즘 설명")
    for item in _algorithm_operation_lines(search_result):
        lines.append(item)
    lines.append("")
    lines.append("결과 요약")
    if first_failure:
        lines.append(
            f"최초 허용치 이탈: 증가값 {float(first_failure['sweep_scale']):.2f} ({float(first_failure['sweep_percent']):.1f}%), 최소전압 {float(first_failure['min_voltage']):.4f} p.u., 최대전압 {float(first_failure['max_voltage']):.4f} p.u., 최대선로 {float(first_failure['max_line_mva']):.3f} MVA"
        )
        fv = first_failure.get("first_violation") or {}
        if fv:
            lines.append(
                f"최초 위반 시각 {fv.get('time', '-')}, 원인 {fv.get('cause', '-')}, 저전압 버스 {fv.get('min_bus_name', '-')}, 과전압 버스 {fv.get('max_bus_name', '-')}, 선로 {fv.get('line_name', '-')}"
            )
    else:
        lines.append("설정한 증가 범위 내에서는 전압 및 선로용량 허용치 이탈이 발생하지 않았습니다.")

    lines.append("")
    lines.append("증가 단계별 판정")
    for idx, run in enumerate(runs):
        lines.append(
            f"증가값 {float(run['sweep_scale']):.2f} ({float(run['sweep_percent']):.1f}%): Vmin {float(run['min_voltage']):.4f}, Vmax {float(run['max_voltage']):.4f}, Line {float(run['max_line_mva']):.3f} MVA, 전압 {'적합' if run['voltage_ok'] else '부적합'}, 선로 {'적합' if run['line_ok'] else '부적합'}"
        )
        lines.append(
            f"운전범위: Load {float(run.get('load_total_min_mw', np.nan)):.2f}~{float(run.get('load_total_max_mw', np.nan)):.2f} MW, PV {float(run.get('pv_total_min_mw', np.nan)):.2f}~{float(run.get('pv_total_max_mw', np.nan)):.2f} MW, WT {float(run.get('wind_total_min_mw', np.nan)):.2f}~{float(run.get('wind_total_max_mw', np.nan)):.2f} MW, ESS {float(run.get('ess_power_min_mw', np.nan)):.2f}~{float(run.get('ess_power_max_mw', np.nan)):.2f} MW, SOC {float(run.get('ess_soc_min_pct', np.nan)):.2f}~{float(run.get('ess_soc_max_pct', np.nan)):.2f}%, OLTC {float(run.get('oltc_tap_min', np.nan)):.0f}~{float(run.get('oltc_tap_max', np.nan)):.0f} tap"
        )
        if search_result.get("scenario") == SCENARIO_RENEWABLE_BY_LOAD_LEVEL:
            lines.append(
                f"부하구간별 최소전압: 경부하 {float(run.get('경부하_min_voltage', np.nan)):.4f}, 중간부하 {float(run.get('중간부하_min_voltage', np.nan)):.4f}, 중부하 {float(run.get('중부하_min_voltage', np.nan)):.4f}"
            )
        if idx < len(run_details):
            for item in _run_evaluation_lines(run_details[idx], search_result):
                lines.append(item)

    lines.append("")
    lines.append("제어기 동작 요약")
    lines.append(f"OLTC 동작 횟수: {int(last_events.get('oltc_moves', 0))}")
    lines.append(f"최종 탭 위치: {int(last_events.get('final_tap', 0))}")
    lines.append(f"ESS 충전 에너지: {float(last_events.get('ess_charge_mwh', 0.0)):.3f} MWh")
    lines.append(f"ESS 방전 에너지: {float(last_events.get('ess_discharge_mwh', 0.0)):.3f} MWh")
    lines.append(f"ESS SOC 범위: {float(last_events.get('soc_min', np.nan)):.2f}% ~ {float(last_events.get('soc_max', np.nan)):.2f}%")

    lines.append("")
    lines.append("해석 및 제언")
    if any(not bool(run["line_ok"]) for run in runs):
        lines.append("선로용량 12 MVA 제약이 먼저 위반되는 구간이 존재합니다. 과부하 선로 인근 ESS 출력을 우선 배치하고 OLTC는 전압 보정 위주로 운전하는 것이 적절합니다.")
    if any(float(run["min_voltage"]) < float(limits.get("voltage_min", 0.94)) for run in runs):
        lines.append("저전압 구간이 존재합니다. OLTC 지연시간을 단축하거나 ESS 방전 개입 이득을 상향하는 것이 유효합니다.")
    if any(float(run["max_voltage"]) > float(limits.get("voltage_max", 1.06)) for run in runs):
        lines.append("과전압 구간이 존재합니다. 재생에너지 증가 구간에서 OLTC 탭다운과 ESS 충전 우선순위를 강화할 필요가 있습니다.")
    if float(last_events.get("soc_max", 0.0)) >= 95.0 or float(last_events.get("soc_min", 100.0)) <= 5.0:
        lines.append("ESS SOC 포화가 관찰됩니다. ESS 용량 또는 운전 목표 SOC 범위를 조정해야 합니다.")
    if search_result.get("scenario") == SCENARIO_RENEWABLE_BY_LOAD_LEVEL:
        lines.append("부하구간별 분석에서는 동일 재생에너지 증가라도 경부하 시 과전압, 중부하 시 저전압 또는 선로혼잡이 다르게 나타날 수 있으므로 운영 기준을 구간별로 분리하는 것이 타당합니다.")
    lines.append("현재 구현은 pandapower AC 조류해석을 유지한 상태기반 협조제어입니다. DistFlow 수식이 추가로 제공되면 민감도 계산 항을 해당 식 기반으로 대체할 수 있습니다.")
    lines.append("")
    lines.append("이상 결과 원인 분석")
    for item in _abnormal_diagnosis(search_result):
        lines.append(item)
    return lines



def _abnormal_diagnosis(search_result: Dict[str, Any]) -> list[str]:
    runs = search_result.get("runs", [])
    first_failure = search_result.get("first_failure") or {}
    last_events = search_result.get("last_events") or {}
    run_details = search_result.get("run_details", [])
    scenario = search_result.get("scenario")
    line_limit = float(search_result.get("limits", {}).get("line_limit_mva", 12.0))
    if not first_failure:
        return ["설정한 범위 내에서는 비정상 결과가 관찰되지 않았습니다. 현재 협조제어 파라미터는 주어진 조건에서 허용치 내 성능을 보였습니다."]

    diagnosis = []
    violation = first_failure.get("first_violation") or {}
    cause = str(violation.get("cause", ""))
    min_v = float(first_failure.get("min_voltage", np.nan))
    max_v = float(first_failure.get("max_voltage", np.nan))
    max_line = float(first_failure.get("max_line_mva", np.nan))

    if ("선로용량" in cause or "line" in cause) and max_line > line_limit:
        diagnosis.append(f"첫 이탈 원인은 선로용량 제약입니다. 전압보다 선로 조류가 먼저 한계 {line_limit:.2f} MVA를 넘었으므로 ESS 위치 또는 선로 혼잡 완화 이득이 부족했을 가능성이 큽니다.")
    if any((not bool(run.get("line_ok", True))) and bool(run.get("voltage_ok", False)) for run in runs):
        diagnosis.append("전압이 허용범위 안에 있어도 선로용량은 별도로 초과될 수 있습니다. 현재 결과는 강한 계통 전원과 짧은 선로로 인해 전압은 비교적 유지되지만, 역송전 전력이 커져 열적 한계가 먼저 위반된 상황에 가깝습니다.")
    voltage_min = float(search_result.get("limits", {}).get("voltage_min", 0.94))
    voltage_max = float(search_result.get("limits", {}).get("voltage_max", 1.06))
    if ("전압" in cause or "voltage" in cause) and min_v < voltage_min:
        diagnosis.append(f"첫 이탈 원인은 저전압입니다. 부하 증가 또는 말단 전압강하에 비해 OLTC 탭업과 ESS 방전 보조가 충분히 빠르지 않아 전압 하한 {voltage_min:.3f} p.u.를 밑돌았습니다.")
    if ("전압" in cause or "voltage" in cause) and max_v > voltage_max:
        diagnosis.append(f"첫 이탈 원인은 과전압입니다. 재생에너지 증가 구간에서 OLTC 탭다운 또는 ESS 충전 흡수가 부족해 전압 상한 {voltage_max:.3f} p.u.를 넘었습니다.")
    if float(last_events.get("soc_min", 100.0)) <= 5.0:
        diagnosis.append("ESS SOC가 하한에 가까워졌습니다. 에너지 용량 부족 또는 방전 전략 과개입 때문에 후반부 제어 여력이 줄었을 수 있습니다.")
    if float(last_events.get("soc_max", 0.0)) >= 95.0:
        diagnosis.append("ESS SOC가 상한에 가까워졌습니다. 재생에너지 흡수 여력이 부족해 과전압 억제 성능이 떨어졌을 가능성이 큽니다.")
    if int(last_events.get("oltc_moves", 0)) >= 20:
        diagnosis.append("OLTC 동작 횟수가 많습니다. 현재 지연시간 또는 전압 deadband가 좁아 탭 동작 빈도가 높게 나타난 것으로 보입니다.")
    if scenario == SCENARIO_RENEWABLE_BY_LOAD_LEVEL:
        diagnosis.append("부하구간별 분석에서는 같은 재생에너지 증가라도 경부하에서는 과전압, 중부하에서는 저전압·혼잡이 다르게 발생할 수 있습니다. 구간별 제어 기준을 분리하는 것이 바람직합니다.")
    if scenario == SCENARIO_LOAD_INCREASE:
        diagnosis.append("부하 증가 시나리오에서는 말단 전압강하와 선로조류 증가가 동시에 발생합니다. ESS를 말단에 배치하더라도 출력 한계가 작으면 저전압과 혼잡을 동시에 막기 어렵습니다.")
    if scenario in [SCENARIO_RENEWABLE_INCREASE, SCENARIO_RENEWABLE_BY_LOAD_LEVEL]:
        diagnosis.append("재생에너지 증가 시나리오에서는 낮은 부하 시간대의 전압 상승이 핵심 원인입니다. ESS 충전 개입을 더 이르게 하거나 OLTC 탭다운 민감도를 높이는 것이 유효합니다.")
    if scenario == SCENARIO_BOTH_INCREASE:
        diagnosis.append("부하와 재생에너지 동시 증가 시나리오에서는 시간대에 따라 저전압과 과전압, 혼잡이 번갈아 나타날 수 있습니다. 따라서 OLTC deadband와 ESS 충방전 전환 조건을 함께 조정하는 것이 중요합니다.")
    return diagnosis

def _render_topology_png(bus_df: pd.DataFrame) -> bytes:
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(10, 2.2))
    ax.axis("off")
    xs = [0.05, 0.23, 0.41, 0.59, 0.77, 0.95]
    labels = ["Substation\nOLTC"]
    for _, row in bus_df.iterrows():
        comps = []
        if float(row[LOAD_P_COL]) > 0:
            comps.append("Load")
        if float(row[PV_P_COL]) > 0:
            comps.append("PV")
        if float(row[WIND_P_COL]) > 0:
            comps.append("Wind")
        if float(row[ESS_MAX_COL]) > 0:
            comps.append("ESS")
        labels.append(f"{row.iloc[0]}\n" + ", ".join(comps))
    for idx, label in enumerate(labels):
        ax.text(xs[idx], 0.5, label, ha="center", va="center", fontsize=10, bbox=dict(boxstyle="round,pad=0.35", fc="#f2f5f8", ec="#5f6b7a"))
        if idx < len(labels) - 1:
            ax.annotate("", xy=(xs[idx + 1] - 0.06, 0.5), xytext=(xs[idx] + 0.06, 0.5), arrowprops=dict(arrowstyle="->", lw=1.3, color="#444"))
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _render_profile_png(last_results: Dict[str, pd.DataFrame]) -> Optional[bytes]:
    df_profile = last_results.get("df_profile", pd.DataFrame()) if last_results else pd.DataFrame()
    if df_profile.empty:
        return None
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(8.5, 3.2))
    x = range(len(df_profile))
    ax.plot(x, df_profile["effective_load_pct"], label="Load %")
    ax.plot(x, df_profile["effective_pv_pct"], label="PV %")
    ax.plot(x, df_profile["effective_wind_pct"], label="Wind %")
    ax.set_title("Load and Renewable Profiles")
    ax.set_xlabel("Time Step")
    ax.set_ylabel("Percent")
    ax.grid(True, alpha=0.3)
    ax.legend(loc="upper right")
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _render_sensitivity_png(search_result: Dict[str, Any]) -> Optional[bytes]:
    runs_df = pd.DataFrame(search_result.get("runs", []))
    if runs_df.empty:
        return None
    import matplotlib.pyplot as plt

    fig, ax1 = plt.subplots(figsize=(8.5, 3.4))
    ax1.plot(runs_df["sweep_percent"], runs_df["min_voltage"], marker="o", label="Min Voltage")
    ax1.plot(runs_df["sweep_percent"], runs_df["max_voltage"], marker="o", label="Max Voltage")
    ax1.set_xlabel("Sweep Percent")
    ax1.set_ylabel("Voltage (p.u.)")
    ax1.axhline(float(search_result.get("limits", {}).get("voltage_min", 0.94)), linestyle="--", color="red", linewidth=1)
    ax1.axhline(float(search_result.get("limits", {}).get("voltage_max", 1.06)), linestyle="--", color="red", linewidth=1)
    ax1.grid(True, alpha=0.3)
    ax2 = ax1.twinx()
    ax2.plot(runs_df["sweep_percent"], runs_df["max_line_mva"], marker="s", color="tab:green", label="Max Line MVA")
    ax2.set_ylabel("Line MVA")
    fig.suptitle("Sensitivity Results")
    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1 + lines2, labels1 + labels2, loc="upper left")
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _build_rich_docx_bytes(search_result: Dict[str, Any]) -> bytes:
    from docx import Document
    from docx.shared import Inches

    document = Document()
    _add_page_number_footer(document)
    document.add_heading("OLTC-ESS Coordinated Control Report", level=0)

    config = search_result.get("config", {})
    bus_df = search_result.get("bus_df", pd.DataFrame())
    last_results = search_result.get("last_results") or {}
    last_events = search_result.get("last_events") or {}
    run_details = search_result.get("run_details", [])

    document.add_heading("1. Base Settings", level=1)
    document.add_paragraph(f"Scenario: {search_result.get('scenario_label', search_result.get('scenario', '-'))}")
    document.add_paragraph(f"Time Step: {int(search_result.get('time_step_mins', 10))} min")
    document.add_paragraph(f"Voltage Limit: {float(search_result.get('limits', {}).get('voltage_min', 0.94)):.3f} ~ {float(search_result.get('limits', {}).get('voltage_max', 1.06)):.3f} p.u.")
    document.add_paragraph(f"Line Limit: {float(search_result.get('limits', {}).get('line_limit_mva', 12.0)):.2f} MVA")
    document.add_paragraph(f"ESS Location / Rating: Bus {int(config.get('ess_bus_number', 5))} / {float(config.get('ess_power_mw', 5.0)):.2f} MW / {float(config.get('ess_capacity_mwh', 15.0)):.2f} MWh")

    document.add_heading("2. Network Topology and Components", level=1)
    topo_png = _render_topology_png(bus_df)
    document.add_picture(io.BytesIO(topo_png), width=Inches(6.8))
    table = document.add_table(rows=1, cols=5)
    headers = ["Bus", "Load (MW)", "PV (MW)", "Wind (MW)", "ESS (MW)"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    if isinstance(bus_df, pd.DataFrame) and not bus_df.empty:
        for _, row in bus_df.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row.iloc[0])
            cells[1].text = f"{float(row[LOAD_P_COL]):.2f}"
            cells[2].text = f"{float(row[PV_P_COL]):.2f}"
            cells[3].text = f"{float(row[WIND_P_COL]):.2f}"
            cells[4].text = f"{float(row[ESS_MAX_COL]):.2f}"

    document.add_heading("3. Load and Renewable Profiles", level=1)
    profile_png = _render_profile_png(last_results)
    if profile_png is not None:
        document.add_picture(io.BytesIO(profile_png), width=Inches(6.8))

    document.add_heading("4. Coordinated Control Settings", level=1)
    document.add_paragraph(f"OLTC Tap Step: {float(config.get('oltc_step', 1.25)):.2f}%")
    document.add_paragraph(f"OLTC Delay / Return Delay: {float(config.get('oltc_delay_mins', 3.0)):.1f} / {float(config.get('oltc_return_delay_mins', 20.0)):.1f} min")
    document.add_paragraph(f"ESS P / Q Gain: {float(config.get('ess_p_gain', 16.0)):.2f} / {float(config.get('ess_q_gain', 8.0)):.2f}")
    document.add_paragraph(f"ESS Ramp Rate: {float(config.get('ess_ramp_rate_mw_per_min', 0.2)):.2f} MW/min")
    document.add_paragraph("Priority: line congestion first, voltage support second. OLTC leads voltage correction and ESS supports fast balancing.")
    for line in _algorithm_operation_lines(search_result):
        document.add_paragraph(line)

    document.add_heading("5. Sensitivity Overview", level=1)
    sens_png = _render_sensitivity_png(search_result)
    if sens_png is not None:
        document.add_picture(io.BytesIO(sens_png), width=Inches(6.8))

    document.add_heading("6. Run-by-Run Results", level=1)
    for run_detail in run_details:
        document.add_heading(_run_detail_title_en(run_detail), level=2)
        for line in _run_detail_lines(run_detail):
            document.add_paragraph(line)
        for line in _run_evaluation_lines(run_detail, search_result):
            document.add_paragraph(line)
        run_png = _render_run_detail_png(run_detail, search_result)
        if run_png is not None:
            document.add_picture(io.BytesIO(run_png), width=Inches(6.6))
        tap_png = _render_oltc_timeseries_png(run_detail.get('results', {}).get('df_tap', pd.DataFrame()), f"OLTC Tap Position: {_run_detail_title_en(run_detail)}")
        if tap_png is not None:
            document.add_picture(io.BytesIO(tap_png), width=Inches(6.6))
        ess_png = _render_ess_power_timeseries_png(run_detail.get('results', {}).get('df_ess_p', pd.DataFrame()), search_result, f"ESS Charge / Discharge: {_run_detail_title_en(run_detail)}")
        if ess_png is not None:
            document.add_picture(io.BytesIO(ess_png), width=Inches(6.6))

    document.add_heading("7. Combined Overlay Graphs", level=1)
    overlay_voltage_png = _render_overlay_voltage_png(run_details, search_result)
    if overlay_voltage_png is not None:
        document.add_picture(io.BytesIO(overlay_voltage_png), width=Inches(6.8))
    overlay_line_png = _render_overlay_line_png(run_details, search_result)
    if overlay_line_png is not None:
        document.add_picture(io.BytesIO(overlay_line_png), width=Inches(6.8))
    overlay_oltc_png = _render_overlay_oltc_png(run_details)
    if overlay_oltc_png is not None:
        document.add_picture(io.BytesIO(overlay_oltc_png), width=Inches(6.8))
    overlay_ess_png = _render_overlay_ess_png(run_details, search_result)
    if overlay_ess_png is not None:
        document.add_picture(io.BytesIO(overlay_ess_png), width=Inches(6.8))

    document.add_heading("8. Operating Range Summary", level=1)
    runs_df = pd.DataFrame(search_result.get("runs", []))
    if not runs_df.empty:
        summary_table = document.add_table(rows=1, cols=8)
        headers = ["Sweep%", "Vmin", "Vmax", "Line", "Load(min~max)", "PV(min~max)", "WT(min~max)", "OLTC(min~max)"]
        for idx, header in enumerate(headers):
            summary_table.rows[0].cells[idx].text = header
        for _, run in runs_df.iterrows():
            cells = summary_table.add_row().cells
            cells[0].text = f"{float(run['sweep_percent']):.1f}"
            cells[1].text = f"{float(run['min_voltage']):.4f}"
            cells[2].text = f"{float(run['max_voltage']):.4f}"
            cells[3].text = f"{float(run['max_line_mva']):.3f}"
            cells[4].text = f"{float(run.get('load_total_min_mw', np.nan)):.2f}~{float(run.get('load_total_max_mw', np.nan)):.2f}"
            cells[5].text = f"{float(run.get('pv_total_min_mw', np.nan)):.2f}~{float(run.get('pv_total_max_mw', np.nan)):.2f}"
            cells[6].text = f"{float(run.get('wind_total_min_mw', np.nan)):.2f}~{float(run.get('wind_total_max_mw', np.nan)):.2f}"
            cells[7].text = f"{float(run.get('oltc_tap_min', np.nan)):.0f}~{float(run.get('oltc_tap_max', np.nan)):.0f}"

    document.add_heading("9. Interpretation and Recommendations", level=1)
    for line in _report_lines(search_result):
        if line:
            document.add_paragraph(line)

    document.add_heading("10. Final State", level=1)
    document.add_paragraph(f"Final Min / Max Voltage: {float(last_events.get('global_min_voltage', np.nan)):.4f} / {float(last_events.get('global_max_voltage', np.nan)):.4f} p.u.")
    document.add_paragraph(f"Maximum Line Loading: {float(last_events.get('global_max_line_mva', np.nan)):.3f} MVA")
    document.add_paragraph(f"OLTC Operations / Final Tap: {int(last_events.get('oltc_moves', 0))} / {int(last_events.get('final_tap', 0))}")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()

def _run_range_report_lines(search_result: Dict[str, Any]) -> list[str]:
    runs_df = pd.DataFrame(search_result.get("runs", []))
    if runs_df.empty:
        return ["시뮬레이션별 운영 범위 데이터가 없습니다."]
    lines = []
    for _, run in runs_df.iterrows():
        lines.append(
            f"{float(run['sweep_percent']):.1f}% | V {float(run['min_voltage']):.4f}~{float(run['max_voltage']):.4f} p.u. | "
            f"Line {float(run['max_line_mva']):.3f} MVA | Load {float(run.get('load_total_min_mw', np.nan)):.2f}~{float(run.get('load_total_max_mw', np.nan)):.2f} MW | "
            f"PV {float(run.get('pv_total_min_mw', np.nan)):.2f}~{float(run.get('pv_total_max_mw', np.nan)):.2f} MW | "
            f"WT {float(run.get('wind_total_min_mw', np.nan)):.2f}~{float(run.get('wind_total_max_mw', np.nan)):.2f} MW | "
            f"ESS P {float(run.get('ess_power_min_mw', np.nan)):.2f}~{float(run.get('ess_power_max_mw', np.nan)):.2f} MW | "
            f"ESS SOC {float(run.get('ess_soc_min_pct', np.nan)):.1f}~{float(run.get('ess_soc_max_pct', np.nan)):.1f}% | "
            f"OLTC Tap {float(run.get('oltc_tap_min', np.nan)):.0f}~{float(run.get('oltc_tap_max', np.nan)):.0f}"
        )
    return lines


def _run_detail_title(run_detail: Dict[str, Any]) -> str:
    return (
        f"증가 배율 {float(run_detail.get('sweep_percent', np.nan)):.1f}% "
        f"(Load {float(run_detail.get('load_scale', np.nan)) * 100:.1f}%, Renewable {float(run_detail.get('renewable_scale', np.nan)) * 100:.1f}%)"
    )


def _algorithm_operation_lines(search_result: Dict[str, Any]) -> list[str]:
    config = search_result.get("config", {})
    scenario = search_result.get("scenario")
    lines = [
        "현재 운영 알고리즘은 상태기반 협조제어로 동작하며, 선로 혼잡을 전압 이상보다 우선 처리합니다.",
        f"OLTC는 ±8 tap, 탭당 {float(config.get('oltc_step', 1.25)):.2f}%, 동작 지연 {float(config.get('oltc_delay_mins', 3.0)):.1f}분, 복귀 지연 {float(config.get('oltc_return_delay_mins', 20.0)):.1f}분 조건으로 운전됩니다.",
        f"ESS는 유효전력/무효전력 보조 제어를 수행하며, P 이득 {float(config.get('ess_p_gain', 16.0)):.2f}, Q 이득 {float(config.get('ess_q_gain', 8.0)):.2f}, 선로 혼잡 완화 이득 {float(config.get('line_relief_gain', 3.0)):.2f}, SOC lock {float(config.get('ess_min_soc', 10.0)):.1f}%~{float(config.get('ess_max_soc', 90.0)):.1f}%를 사용합니다.",
    ]
    if scenario == SCENARIO_LOAD_INCREASE:
        lines.append("부하 증가 시나리오에서는 OLTC 탭업과 ESS 방전이 전압 유지와 선로용량 확보의 핵심 대응입니다.")
    elif scenario == SCENARIO_RENEWABLE_INCREASE:
        lines.append("재생에너지 증가 시나리오에서는 OLTC 탭다운과 ESS 충전이 과전압 및 역조류 완화의 핵심 대응입니다.")
    elif scenario == SCENARIO_BOTH_INCREASE:
        lines.append("부하와 재생에너지 동시 증가 시나리오에서는 시간대에 따라 탭업/탭다운, ESS 충전/방전이 교차할 수 있어 두 제어기의 협조 동작이 더 중요합니다.")
    elif scenario == SCENARIO_RENEWABLE_BY_LOAD_LEVEL:
        lines.append("부하구간별 재생에너지 증가 시나리오에서는 경부하 시간대 과전압과 중부하 시간대 선로혼잡 또는 저전압을 구간별로 다르게 대응합니다.")
    return lines


def _run_detail_lines(run_detail: Dict[str, Any]) -> list[str]:
    verdict = "적합" if bool(run_detail.get("overall_ok", False)) else "부적합"
    return [
        f"판정: {verdict} | 최소/최대 전압 {float(run_detail.get('min_voltage', np.nan)):.4f} ~ {float(run_detail.get('max_voltage', np.nan)):.4f} p.u. | 최대 선로용량 {float(run_detail.get('max_line_mva', np.nan)):.3f} MVA",
        f"부하 총합 {float(run_detail.get('load_total_min_mw', np.nan)):.2f} ~ {float(run_detail.get('load_total_max_mw', np.nan)):.2f} MW | PV 총합 {float(run_detail.get('pv_total_min_mw', np.nan)):.2f} ~ {float(run_detail.get('pv_total_max_mw', np.nan)):.2f} MW | WT 총합 {float(run_detail.get('wind_total_min_mw', np.nan)):.2f} ~ {float(run_detail.get('wind_total_max_mw', np.nan)):.2f} MW",
        f"OLTC 동작 {int(run_detail.get('oltc_moves', 0))}회, 탭 범위 {float(run_detail.get('oltc_tap_min', np.nan)):.0f} ~ {float(run_detail.get('oltc_tap_max', np.nan)):.0f}, 최종 탭 {int(run_detail.get('final_tap', 0))} | ESS 충전 {float(run_detail.get('ess_charge_mwh', 0.0)):.3f} MWh, 방전 {float(run_detail.get('ess_discharge_mwh', 0.0)):.3f} MWh, 출력범위 {float(run_detail.get('ess_power_min_mw', np.nan)):.2f} ~ {float(run_detail.get('ess_power_max_mw', np.nan)):.2f} MW | SOC {float(run_detail.get('ess_soc_min_pct', np.nan)):.1f} ~ {float(run_detail.get('ess_soc_max_pct', np.nan)):.1f}%",
    ]


def _run_evaluation_lines(run_detail: Dict[str, Any], search_result: Dict[str, Any]) -> list[str]:
    config = search_result.get("config", {})
    scenario = search_result.get("scenario")
    ess_rating = max(float(config.get("ess_power_mw", 0.0)), 1e-9)
    ess_util = max(abs(float(run_detail.get("ess_power_min_mw", 0.0))), abs(float(run_detail.get("ess_power_max_mw", 0.0)))) / ess_rating
    tap_span = max(abs(float(run_detail.get("oltc_tap_min", 0.0))), abs(float(run_detail.get("oltc_tap_max", 0.0))))
    voltage_ok = bool(run_detail.get("voltage_ok", False))
    line_ok = bool(run_detail.get("line_ok", False))
    overall_ok = bool(run_detail.get("overall_ok", False))

    mode_text = ""
    if scenario == SCENARIO_LOAD_INCREASE:
        mode_text = "부하 증가 조건에서는 탭업과 ESS 방전이 중심이 되어 저전압과 선로부하 증가를 완화했는지 보는 것이 핵심입니다."
    elif scenario == SCENARIO_RENEWABLE_INCREASE:
        mode_text = "재생에너지 증가 조건에서는 탭다운과 ESS 충전이 과전압 및 역조류 완화에 얼마나 기여했는지 보는 것이 핵심입니다."
    elif scenario == SCENARIO_BOTH_INCREASE:
        mode_text = "부하와 재생에너지가 함께 증가하는 조건에서는 시간대별로 충전/방전과 탭업/탭다운이 교차하며 두 제어기가 균형 있게 작동했는지 보는 것이 핵심입니다."
    elif scenario == SCENARIO_RENEWABLE_BY_LOAD_LEVEL:
        mode_text = "부하구간별 조건에서는 경부하의 과전압 대응과 중부하·중간부하의 혼잡/저전압 대응이 모두 적절했는지 보는 것이 핵심입니다."

    if overall_ok:
        if int(run_detail.get("oltc_moves", 0)) > 0 or float(run_detail.get("ess_charge_mwh", 0.0)) > 0.01 or float(run_detail.get("ess_discharge_mwh", 0.0)) > 0.01:
            assess = "평가: OLTC와 ESS가 실제로 개입해 전압 및 선로용량 유지에 기여를 잘 수행했습니다."
        else:
            assess = "평가: 허용치 내였으며 제어기 개입은 크지 않았습니다. 계통 자체 여유가 충분했던 것으로 해석됩니다."
    else:
        reasons = []
        if not voltage_ok:
            reasons.append("전압 유지")
        if not line_ok:
            reasons.append("선로용량 유지")
        if ess_util >= 0.95 or float(run_detail.get("ess_soc_min_pct", 100.0)) <= float(config.get("ess_min_soc", 10.0)) + 1.0 or float(run_detail.get("ess_soc_max_pct", 0.0)) >= float(config.get("ess_max_soc", 90.0)) - 1.0:
            reasons.append("ESS 여유 부족")
        if tap_span >= 8 or int(run_detail.get("oltc_moves", 0)) >= 16:
            reasons.append("OLTC 여유 또는 동작 빈도 한계")
        reason_text = ", ".join(reasons) if reasons else "제어 여유 부족"
        assess = f"평가: {reason_text} 때문에 전압 및 선로용량 유지에 충분히 기여하지 못했습니다."

    return [mode_text, assess]


def _apply_time_axis(ax, labels: list[str]):
    if not labels:
        return
    tick_step = max(1, len(labels) // 8)
    tick_idx = list(range(0, len(labels), tick_step))
    if tick_idx[-1] != len(labels) - 1:
        tick_idx.append(len(labels) - 1)
    ax.set_xticks(tick_idx)
    ax.set_xticklabels([labels[idx] for idx in tick_idx], rotation=45, ha="right")


def _render_voltage_line_timeseries_png(df_min: pd.DataFrame, df_max: pd.DataFrame, df_line: pd.DataFrame, search_result: Dict[str, Any], title: str) -> Optional[bytes]:
    if df_min.empty or df_max.empty or df_line.empty:
        return None

    import matplotlib.pyplot as plt

    labels = [str(idx) for idx in df_min.index]
    x = list(range(len(labels)))

    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(8.5, 5.0), sharex=True)
    ax1.plot(x, df_min.iloc[:, 0], label="Min Voltage", color="tab:blue")
    ax1.plot(x, df_max.iloc[:, 0], label="Max Voltage", color="tab:orange")
    ax1.axhline(float(search_result.get("limits", {}).get("voltage_min", 0.94)), linestyle="--", color="red", linewidth=1)
    ax1.axhline(float(search_result.get("limits", {}).get("voltage_max", 1.06)), linestyle="--", color="red", linewidth=1)
    ax1.set_ylabel("Voltage (p.u.)")
    ax1.set_title(title)
    ax1.grid(True, alpha=0.3)
    ax1.legend(loc="upper right")

    ax2.plot(x, df_line.iloc[:, 0], label="Max Line MVA", color="tab:green")
    ax2.axhline(float(search_result.get("limits", {}).get("line_limit_mva", 12.0)), linestyle="--", color="red", linewidth=1)
    ax2.set_xlabel("Time")
    ax2.set_ylabel("MVA")
    ax2.grid(True, alpha=0.3)
    ax2.legend(loc="upper right")
    _apply_time_axis(ax2, labels)

    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _render_representative_timeseries_png(results: Dict[str, pd.DataFrame], search_result: Dict[str, Any]) -> Optional[bytes]:
    df_min = results.get("df_min_v", pd.DataFrame()) if results else pd.DataFrame()
    df_max = results.get("df_max_v", pd.DataFrame()) if results else pd.DataFrame()
    df_line = results.get("df_line_mva_max", pd.DataFrame()) if results else pd.DataFrame()
    return _render_voltage_line_timeseries_png(df_min, df_max, df_line, search_result, "Representative Run: Voltage and Line Loading by Time")


def _run_detail_title_en(run_detail: Dict[str, Any]) -> str:
    return (
        f"Run {float(run_detail.get('sweep_percent', np.nan)):.1f}% "
        f"(Load {float(run_detail.get('load_scale', np.nan)) * 100:.1f}%, Renewable {float(run_detail.get('renewable_scale', np.nan)) * 100:.1f}%)"
    )


def _active_ess_series(df_ess_p: pd.DataFrame, search_result: Dict[str, Any]) -> tuple[Optional[str], pd.Series]:
    if df_ess_p.empty:
        return None, pd.Series(dtype=float)
    config = search_result.get("config", {})
    preferred = f"Bus {int(config.get('ess_bus_number', 5))}"
    if preferred in df_ess_p.columns:
        return preferred, pd.to_numeric(df_ess_p[preferred], errors="coerce")
    best_col = None
    best_mag = -1.0
    for col in df_ess_p.columns:
        series = pd.to_numeric(df_ess_p[col], errors="coerce")
        mag = float(series.abs().max(skipna=True)) if not series.empty else 0.0
        if mag > best_mag:
            best_mag = mag
            best_col = col
    if best_col is None:
        return None, pd.Series(dtype=float)
    return best_col, pd.to_numeric(df_ess_p[best_col], errors="coerce")


def _render_run_detail_png(run_detail: Dict[str, Any], search_result: Dict[str, Any]) -> Optional[bytes]:
    detail_results = run_detail.get("results", {}) if run_detail else {}
    df_min = detail_results.get("df_min_v", pd.DataFrame())
    df_max = detail_results.get("df_max_v", pd.DataFrame())
    df_line = detail_results.get("df_line_mva_max", pd.DataFrame())
    return _render_voltage_line_timeseries_png(df_min, df_max, df_line, search_result, f"Voltage and Line Loading: {_run_detail_title_en(run_detail)}")


def _render_oltc_timeseries_png(df_tap: pd.DataFrame, title: str) -> Optional[bytes]:
    if df_tap.empty or "OLTC Tap" not in df_tap.columns:
        return None
    import matplotlib.pyplot as plt

    labels = [str(idx) for idx in df_tap.index]
    x = list(range(len(labels)))
    fig, ax = plt.subplots(figsize=(8.5, 2.8))
    ax.step(x, df_tap["OLTC Tap"], where="post", label="OLTC Tap Position", color="tab:purple")
    ax.set_title(title)
    ax.set_xlabel("Time")
    ax.set_ylabel("Tap Position")
    ax.grid(True, alpha=0.3)
    ax.legend(loc="upper right")
    _apply_time_axis(ax, labels)

    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _render_ess_power_timeseries_png(df_ess_p: pd.DataFrame, search_result: Dict[str, Any], title: str) -> Optional[bytes]:
    col, series = _active_ess_series(df_ess_p, search_result)
    if col is None or series.empty:
        return None
    import matplotlib.pyplot as plt

    labels = [str(idx) for idx in df_ess_p.index]
    x = list(range(len(labels)))
    fig, ax = plt.subplots(figsize=(8.5, 2.8))
    ax.plot(x, series, label=f"ESS Active Power ({col})", color="tab:cyan")
    ax.axhline(0.0, color="black", linewidth=1.0, linestyle="--")
    ax.set_title(title)
    ax.set_xlabel("Time")
    ax.set_ylabel("Active Power (MW)")
    ax.grid(True, alpha=0.3)
    ax.legend(loc="upper right")
    _apply_time_axis(ax, labels)

    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _render_overlay_voltage_png(run_details: list[Dict[str, Any]], search_result: Dict[str, Any]) -> Optional[bytes]:
    if not run_details:
        return None
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(8.8, 4.2))
    labels = []
    for run_detail in run_details:
        detail_results = run_detail.get("results", {})
        df_min = detail_results.get("df_min_v", pd.DataFrame())
        df_max = detail_results.get("df_max_v", pd.DataFrame())
        if df_min.empty or df_max.empty:
            continue
        labels = [str(idx) for idx in df_min.index]
        x = list(range(len(labels)))
        base = f"{float(run_detail.get('sweep_percent', np.nan)):.1f}%"
        ax.plot(x, df_min.iloc[:, 0], label=f"{base} Min")
        ax.plot(x, df_max.iloc[:, 0], linestyle="--", linewidth=1.0, label=f"{base} Max")
    if not labels:
        plt.close(fig)
        return None
    ax.axhline(float(search_result.get("limits", {}).get("voltage_min", 0.94)), linestyle="--", color="red", linewidth=1)
    ax.axhline(float(search_result.get("limits", {}).get("voltage_max", 1.06)), linestyle="--", color="red", linewidth=1)
    ax.set_title("Combined Voltage Profiles by Run")
    ax.set_xlabel("Time")
    ax.set_ylabel("Voltage (p.u.)")
    ax.grid(True, alpha=0.3)
    ax.legend(loc="upper left", bbox_to_anchor=(1.02, 1.0), fontsize=8)
    _apply_time_axis(ax, labels)

    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _render_overlay_line_png(run_details: list[Dict[str, Any]], search_result: Dict[str, Any]) -> Optional[bytes]:
    if not run_details:
        return None
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(8.8, 4.0))
    labels = []
    for run_detail in run_details:
        df_line = run_detail.get("results", {}).get("df_line_mva_max", pd.DataFrame())
        if df_line.empty:
            continue
        labels = [str(idx) for idx in df_line.index]
        x = list(range(len(labels)))
        base = f"{float(run_detail.get('sweep_percent', np.nan)):.1f}%"
        ax.plot(x, df_line.iloc[:, 0], label=base)
    if not labels:
        plt.close(fig)
        return None
    ax.axhline(float(search_result.get("limits", {}).get("line_limit_mva", 12.0)), linestyle="--", color="red", linewidth=1)
    ax.set_title("Combined Line Loading by Run")
    ax.set_xlabel("Time")
    ax.set_ylabel("Line Loading (MVA)")
    ax.grid(True, alpha=0.3)
    ax.legend(loc="upper left", bbox_to_anchor=(1.02, 1.0), fontsize=8)
    _apply_time_axis(ax, labels)

    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _render_overlay_oltc_png(run_details: list[Dict[str, Any]]) -> Optional[bytes]:
    if not run_details:
        return None
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(8.8, 4.0))
    labels = []
    for run_detail in run_details:
        df_tap = run_detail.get("results", {}).get("df_tap", pd.DataFrame())
        if df_tap.empty or "OLTC Tap" not in df_tap.columns:
            continue
        labels = [str(idx) for idx in df_tap.index]
        x = list(range(len(labels)))
        base = f"{float(run_detail.get('sweep_percent', np.nan)):.1f}%"
        ax.step(x, df_tap["OLTC Tap"], where="post", label=base)
    if not labels:
        plt.close(fig)
        return None
    ax.set_title("Combined OLTC Tap Position by Run")
    ax.set_xlabel("Time")
    ax.set_ylabel("Tap Position")
    ax.grid(True, alpha=0.3)
    ax.legend(loc="upper left", bbox_to_anchor=(1.02, 1.0), fontsize=8)
    _apply_time_axis(ax, labels)

    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _render_overlay_ess_png(run_details: list[Dict[str, Any]], search_result: Dict[str, Any]) -> Optional[bytes]:
    if not run_details:
        return None
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(8.8, 4.0))
    labels = []
    for run_detail in run_details:
        df_ess_p = run_detail.get("results", {}).get("df_ess_p", pd.DataFrame())
        col, series = _active_ess_series(df_ess_p, search_result)
        if col is None or series.empty:
            continue
        labels = [str(idx) for idx in df_ess_p.index]
        x = list(range(len(labels)))
        base = f"{float(run_detail.get('sweep_percent', np.nan)):.1f}%"
        ax.plot(x, series, label=base)
    if not labels:
        plt.close(fig)
        return None
    ax.axhline(0.0, color="black", linewidth=1.0, linestyle="--")
    ax.set_title("Combined ESS Charge / Discharge by Run")
    ax.set_xlabel("Time")
    ax.set_ylabel("Active Power (MW)")
    ax.grid(True, alpha=0.3)
    ax.legend(loc="upper left", bbox_to_anchor=(1.02, 1.0), fontsize=8)
    _apply_time_axis(ax, labels)

    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _docx_text_paragraph_xml(text: str, bold: bool = False) -> str:
    safe = escape(str(text))
    run_pr = "<w:rPr><w:b/></w:rPr>" if bold else ""
    return f'<w:p><w:r>{run_pr}<w:t xml:space="preserve">{safe}</w:t></w:r></w:p>'


def _docx_image_paragraph_xml(rel_id: str, drawing_id: int, name: str, cx: int, cy: int) -> str:
    safe_name = escape(name)
    return (
        '<w:p><w:r><w:drawing><wp:inline distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        f'<wp:docPr id="{drawing_id}" name="{safe_name}"/>'
        '<wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>'
        '<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic>'
        f'<pic:nvPicPr><pic:cNvPr id="{drawing_id}" name="{safe_name}"/><pic:cNvPicPr/></pic:nvPicPr>'
        f'<pic:blipFill><a:blip r:embed="{rel_id}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
        f'<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>'
        '</pic:pic></a:graphicData></a:graphic></wp:inline></w:drawing></w:r></w:p>'
    )


def _add_page_number_footer(document) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Page ")

    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _docx_footer_xml() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:ftr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:p><w:pPr><w:jc w:val="center"/></w:pPr>'
        '<w:r><w:t xml:space="preserve">Page </w:t></w:r>'
        '<w:r><w:fldChar w:fldCharType="begin"/></w:r>'
        '<w:r><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>'
        '<w:r><w:fldChar w:fldCharType="end"/></w:r>'
        '</w:p></w:ftr>'
    )


def _build_graph_fallback_docx_bytes(search_result: Dict[str, Any]) -> bytes:
    # python-docx가 없어도 matplotlib 이미지와 수동 OOXML 패키징으로 그래프 포함 Word를 생성한다.
    config = search_result.get("config", {})
    bus_df = search_result.get("bus_df", pd.DataFrame())
    last_results = search_result.get("last_results") or {}
    detail_results = search_result.get("first_failure_results") or search_result.get("last_results") or {}
    run_details = search_result.get("run_details", [])

    body_parts = []
    images: list[tuple[str, str, bytes]] = []

    def add_paragraph(text: str, bold: bool = False):
        if text:
            body_parts.append(_docx_text_paragraph_xml(text, bold=bold))
        else:
            body_parts.append('<w:p/>')

    def add_image_section(title: str, image_bytes: Optional[bytes], width_in: float, height_in: float):
        if not image_bytes:
            return
        index = len(images) + 1
        rel_id = f"rId{index}"
        filename = f"image{index}.png"
        images.append((rel_id, filename, image_bytes))
        add_paragraph(title, bold=True)
        body_parts.append(_docx_image_paragraph_xml(rel_id, index, title, int(width_in * 914400), int(height_in * 914400)))

    add_paragraph("OLTC-ESS Coordinated Control Report", bold=True)
    add_paragraph("")
    add_paragraph("1. Base Settings", bold=True)
    add_paragraph(f"Scenario: {search_result.get('scenario_label', search_result.get('scenario', '-'))}")
    add_paragraph(f"Time Step: {int(search_result.get('time_step_mins', 10))} min")
    add_paragraph(f"Voltage Limit: {float(search_result.get('limits', {}).get('voltage_min', 0.94)):.3f} ~ {float(search_result.get('limits', {}).get('voltage_max', 1.06)):.3f} p.u.")
    add_paragraph(f"Line Limit: {float(search_result.get('limits', {}).get('line_limit_mva', 12.0)):.2f} MVA")
    add_paragraph(f"ESS Location / Rating: Bus {int(config.get('ess_bus_number', 5))} / {float(config.get('ess_power_mw', 5.0)):.2f} MW / {float(config.get('ess_capacity_mwh', 15.0)):.2f} MWh")

    add_paragraph("2. Network Topology and Components", bold=True)
    if isinstance(bus_df, pd.DataFrame) and not bus_df.empty:
        for _, row in bus_df.iterrows():
            add_paragraph(
                f"{row.iloc[0]}: Load {float(row[LOAD_P_COL]):.2f} MW / PV {float(row[PV_P_COL]):.2f} MW / Wind {float(row[WIND_P_COL]):.2f} MW / ESS {float(row[ESS_MAX_COL]):.2f} MW"
            )
    add_image_section("Network Topology", _render_topology_png(bus_df), 6.8, 1.9)

    add_paragraph("3. Load and Renewable Profiles", bold=True)
    add_image_section("Load and Renewable Profiles", _render_profile_png(last_results), 6.8, 2.7)

    add_paragraph("4. Coordinated Control Settings", bold=True)
    add_paragraph(f"OLTC Tap Step: {float(config.get('oltc_step', 1.25)):.2f}%")
    add_paragraph(f"OLTC Delay / Return Delay: {float(config.get('oltc_delay_mins', 3.0)):.1f} / {float(config.get('oltc_return_delay_mins', 20.0)):.1f} min")
    add_paragraph(f"ESS P / Q Gain: {float(config.get('ess_p_gain', 16.0)):.2f} / {float(config.get('ess_q_gain', 8.0)):.2f}")
    add_paragraph(f"ESS Ramp Rate: {float(config.get('ess_ramp_rate_mw_per_min', 0.2)):.2f} MW/min")
    add_paragraph("Priority: line congestion first, voltage support second. OLTC leads voltage correction and ESS supports fast balancing.")
    for line in _algorithm_operation_lines(search_result):
        add_paragraph(line)

    add_paragraph("5. Sensitivity Overview", bold=True)
    add_image_section("Sensitivity Overview", _render_sensitivity_png(search_result), 6.8, 2.9)

    add_paragraph("6. Representative Time-Series", bold=True)
    add_image_section("Representative Voltage and Line Loading", _render_representative_timeseries_png(detail_results, search_result), 6.8, 4.0)

    add_paragraph("7. Run-by-Run Results", bold=True)
    for run_detail in run_details:
        add_paragraph(_run_detail_title_en(run_detail), bold=True)
        for line in _run_detail_lines(run_detail):
            add_paragraph(line)
        for line in _run_evaluation_lines(run_detail, search_result):
            add_paragraph(line)
        add_image_section(f"Voltage and Line Loading: {_run_detail_title_en(run_detail)}", _render_run_detail_png(run_detail, search_result), 6.8, 4.0)
        add_image_section(f"OLTC Tap Position: {_run_detail_title_en(run_detail)}", _render_oltc_timeseries_png(run_detail.get('results', {}).get('df_tap', pd.DataFrame()), f"OLTC Tap Position: {_run_detail_title_en(run_detail)}"), 6.8, 2.8)
        add_image_section(f"ESS Charge / Discharge: {_run_detail_title_en(run_detail)}", _render_ess_power_timeseries_png(run_detail.get('results', {}).get('df_ess_p', pd.DataFrame()), search_result, f"ESS Charge / Discharge: {_run_detail_title_en(run_detail)}"), 6.8, 2.8)

    add_paragraph("8. Combined Overlay Graphs", bold=True)
    add_image_section("Combined Voltage Profiles by Run", _render_overlay_voltage_png(run_details, search_result), 6.8, 4.0)
    add_image_section("Combined Line Loading by Run", _render_overlay_line_png(run_details, search_result), 6.8, 4.0)
    add_image_section("Combined OLTC Tap Position by Run", _render_overlay_oltc_png(run_details), 6.8, 4.0)
    add_image_section("Combined ESS Charge / Discharge by Run", _render_overlay_ess_png(run_details, search_result), 6.8, 4.0)

    add_paragraph("9. Operating Range Summary", bold=True)
    for line in _run_range_report_lines(search_result):
        add_paragraph(line)

    add_paragraph("10. Interpretation and Recommendations", bold=True)
    for line in _report_lines(search_result):
        add_paragraph(line)

    footer_xml = _docx_footer_xml()
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<w:body>' + ''.join(body_parts) + '<w:sectPr><w:footerReference w:type="default" r:id="rIdFooter1"/><w:pgSz w:w="12240" w:h="15840"/>'
        '<w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" w:header="720" w:footer="720" w:gutter="0"/>'
        '</w:sectPr></w:body></w:document>'
    )

    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Default Extension="png" ContentType="image/png"/>'
        '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        '<Override PartName="/word/footer1.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml"/>'
        '</Types>'
    )
    package_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
        '</Relationships>'
    )
    document_rels = ['<?xml version="1.0" encoding="UTF-8" standalone="yes"?>', '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">']
    document_rels.append('<Relationship Id="rIdFooter1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer" Target="footer1.xml"/>')
    for rel_id, filename, _ in images:
        document_rels.append(
            f'<Relationship Id="{rel_id}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/{filename}"/>'
        )
    document_rels.append('</Relationships>')

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", package_rels)
        zf.writestr("word/document.xml", document_xml)
        zf.writestr("word/footer1.xml", footer_xml)
        zf.writestr("word/_rels/document.xml.rels", ''.join(document_rels))
        for _, filename, image_bytes in images:
            zf.writestr(f"word/media/{filename}", image_bytes)
    return buffer.getvalue()

def build_word_report_bytes(search_result: Dict[str, Any]) -> bytes:
    # 1) python-docx rich report 2) 수동 OOXML 그래프 리포트 3) 텍스트형 최소 리포트 순으로 시도한다.
    try:
        return _build_rich_docx_bytes(search_result)
    except Exception as rich_exc:
        try:
            return _build_graph_fallback_docx_bytes(search_result)
        except Exception as graph_exc:
            lines = [
                f"그래프 포함 보고서 생성 실패: {type(rich_exc).__name__} -> {type(graph_exc).__name__}",
                "환경에 python-docx 또는 matplotlib가 누락되었을 가능성이 있습니다.",
                "",
            ] + _report_lines(search_result)
            return _build_docx_bytes("OLTC-ESS 협조제어 분석 보고서", lines)


def write_word_report(report_path: str, search_result: Dict[str, Any]) -> str:
    report_bytes = build_word_report_bytes(search_result)
    with open(report_path, "wb") as f:
        f.write(report_bytes)
    return report_path



































